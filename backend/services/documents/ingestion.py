"""
Document ingestion service.

Responsibilities:
  1. Accept raw file bytes + metadata
  2. Verify SHA-256 checksum (dedup)
  3. Encrypt file bytes with Fernet before persisting to disk
  4. Extract plain text (PDF via PyMuPDF, DOCX via python-docx, .txt/.md direct)
  5. Chunk text (fixed ~512 tokens with 64-token overlap)
  6. Embed each chunk via Ollama nomic-embed-text (768-dim)
  7. Persist Document + DocumentChunk rows in one transaction

Security notes:
  - Raw bytes are NEVER written unencrypted.
  - Fernet key is sourced from OS keyring via settings.fernet_key.
  - File path is stored in Document.source_uri (relative to STORAGE_ROOT).
"""

from __future__ import annotations

import hashlib
import logging
import mimetypes
import re
from pathlib import Path
from uuid import UUID

from cryptography.fernet import Fernet
from sqlalchemy.ext.asyncio import AsyncSession

from config import settings
from models.db.document import Document, DocumentChunk
from repositories.document_repo import DocumentChunkRepository, DocumentRepository
from services.embeddings.service import EmbeddingService

logger = logging.getLogger(__name__)

# ── Constants ─────────────────────────────────────────────────────────────────

CHUNK_TOKENS = 512        # target tokens per chunk
CHUNK_OVERLAP = 64        # token overlap between consecutive chunks
WORDS_PER_TOKEN = 0.75    # rough approximation for splitting by words

SUPPORTED_MIME_TYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "text/plain",
    "text/markdown",
    "text/csv",
}


# ── Text extraction helpers ───────────────────────────────────────────────────

def _extract_text_pdf(raw: bytes) -> str:
    """Extract text from PDF bytes using PyMuPDF (fitz)."""
    try:
        import fitz  # PyMuPDF

        doc = fitz.open(stream=raw, filetype="pdf")
        pages = [page.get_text() for page in doc]
        doc.close()
        return "\n\n".join(pages)
    except ImportError as exc:
        raise RuntimeError("PyMuPDF (pymupdf) is not installed") from exc


def _extract_text_docx(raw: bytes) -> str:
    """Extract text from DOCX bytes using python-docx."""
    try:
        import io

        import docx

        doc = docx.Document(io.BytesIO(raw))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    except ImportError as exc:
        raise RuntimeError("python-docx is not installed") from exc


def _extract_text(raw: bytes, mime_type: str) -> str:
    if mime_type == "application/pdf":
        return _extract_text_pdf(raw)
    if mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return _extract_text_docx(raw)
    # Plain text / markdown / csv
    return raw.decode("utf-8", errors="replace")


# ── Chunking ──────────────────────────────────────────────────────────────────

def _chunk_text(text: str) -> list[str]:
    """
    Split *text* into overlapping chunks of approximately CHUNK_TOKENS tokens.
    Uses whitespace splitting as a token approximation.
    """
    words = text.split()
    if not words:
        return []

    words_per_chunk = int(CHUNK_TOKENS / WORDS_PER_TOKEN)
    overlap_words = int(CHUNK_OVERLAP / WORDS_PER_TOKEN)
    step = words_per_chunk - overlap_words

    chunks: list[str] = []
    start = 0
    while start < len(words):
        end = min(start + words_per_chunk, len(words))
        chunks.append(" ".join(words[start:end]))
        if end == len(words):
            break
        start += step
    return chunks


# ── Storage helpers ───────────────────────────────────────────────────────────

def _get_storage_root() -> Path:
    root = Path(settings.storage_root).expanduser().resolve()
    root.mkdir(parents=True, exist_ok=True)
    return root


def _encrypt_and_store(doc_id: UUID, raw: bytes, extension: str) -> Path:
    """Encrypt *raw* with Fernet and write to STORAGE_ROOT/{doc_id}{ext}.enc."""
    fernet = Fernet(settings.fernet_key)
    encrypted = fernet.encrypt(raw)
    dest = _get_storage_root() / f"{doc_id}{extension}.enc"
    dest.write_bytes(encrypted)
    return dest


def _decrypt_file(doc_id: UUID, extension: str) -> bytes:
    """Read and decrypt a stored document."""
    fernet = Fernet(settings.fernet_key)
    src = _get_storage_root() / f"{doc_id}{extension}.enc"
    return fernet.decrypt(src.read_bytes())


def _delete_stored_file(doc_id: UUID, extension: str) -> None:
    path = _get_storage_root() / f"{doc_id}{extension}.enc"
    if path.exists():
        path.unlink()


async def _ensure_vector_dimension(db: AsyncSession, target_dim: int) -> None:
    """
    Check the current pgvector column dimension; ALTER it if it doesn't match
    *target_dim*.  Mirrors the same logic in the reindex endpoint so a fresh
    ingest after a provider switch doesn't fail with a dimension mismatch.
    """
    from sqlalchemy import text

    row = (await db.execute(text(
        "SELECT atttypmod FROM pg_attribute "
        "JOIN pg_class ON attrelid = pg_class.oid "
        "WHERE relname = 'document_chunks' AND attname = 'embedding'"
    ))).fetchone()
    current_dim = int(row[0]) if row and row[0] else 768

    if current_dim == target_dim:
        return

    logger.info("Auto-altering embedding column: %d → %d dims", current_dim, target_dim)
    await db.execute(text("DROP INDEX IF EXISTS ix_document_chunks_embedding"))
    await db.execute(text(
        f"ALTER TABLE document_chunks "
        f"ALTER COLUMN embedding TYPE vector({target_dim}) USING NULL"
    ))
    # Update the dimension setting
    await db.execute(text(
        f"INSERT INTO system_settings (key, value) VALUES ('llm.embedding_dimension', '{target_dim}') "
        f"ON CONFLICT (key) DO UPDATE SET value = '{target_dim}'"
    ))
    await db.commit()
    logger.info("Column altered to vector(%d) and committed", target_dim)


# ── Main service class ────────────────────────────────────────────────────────

class DocumentIngestionService:

    def __init__(
        self,
        db: AsyncSession,
        embedding_svc: EmbeddingService,
    ) -> None:
        self._db = db
        self._emb = embedding_svc
        self._doc_repo = DocumentRepository(db)
        self._chunk_repo = DocumentChunkRepository(db)

    async def ingest(
        self,
        *,
        filename: str,
        raw_bytes: bytes,
        title: str,
        category_id: UUID | None,
        is_regulated: bool,
        created_by_id: UUID,
    ) -> Document:
        """
        Full ingestion pipeline.  Returns the persisted Document.
        Raises ValueError on unsupported MIME type.
        """
        # 1. Detect MIME type
        mime_type, _ = mimetypes.guess_type(filename)
        if mime_type not in SUPPORTED_MIME_TYPES:
            # Fallback: treat as plain text if extension is .txt/.md
            ext = Path(filename).suffix.lower()
            if ext in (".txt", ".md", ".csv", ".json"):
                mime_type = "text/plain"
            else:
                raise ValueError(
                    f"Unsupported file type: {mime_type or 'unknown'}. "
                    f"Supported: PDF, DOCX, TXT, MD, CSV"
                )

        # 2. Checksum
        checksum = hashlib.sha256(raw_bytes).hexdigest()

        # 3. Persist Document row (status = processing)
        extension = Path(filename).suffix.lower() or ".bin"
        doc = Document(
            category_id=category_id,
            title=title,
            file_name=filename,
            source_type="upload",
            mime_type=mime_type,
            file_size_bytes=len(raw_bytes),
            checksum_sha256=checksum,
            is_regulated=is_regulated,
            status="processing",
            chunk_count=0,
            created_by=created_by_id,
        )
        self._db.add(doc)
        await self._db.flush()   # assigns doc.id

        try:
            # 4. Encrypt and store
            stored_path = _encrypt_and_store(doc.id, raw_bytes, extension)
            doc.source_uri = str(stored_path.relative_to(_get_storage_root()))

            # 5. Extract text
            text = _extract_text(raw_bytes, mime_type)
            text = re.sub(r"\s{3,}", "\n\n", text).strip()

            # 6. Chunk
            chunks = _chunk_text(text)
            if not chunks:
                chunks = [text] if text else ["(no content)"]

            # 7. Embed all chunks in one batch call (1 API request per document,
            #    avoids rate-limit exhaustion on free-tier providers like Voyage).
            #    Use the first embedding to detect dimension and auto-alter column.
            all_embeddings = await self._emb.embed_batch(chunks)
            actual_dim = len(all_embeddings[0])
            await _ensure_vector_dimension(self._db, actual_dim)

            emb_model = getattr(self._emb, "_model", "nomic-embed-text")

            # 8. Create chunk rows
            for idx, (chunk_text, embedding) in enumerate(zip(chunks, all_embeddings)):
                chunk = DocumentChunk(
                    document_id=doc.id,
                    chunk_index=idx,
                    content=chunk_text,
                    content_hash=hashlib.sha256(chunk_text.encode()).hexdigest(),
                    embedding=embedding,
                    embedding_model=emb_model,
                    embedding_dimension=actual_dim,
                    page_number=None,
                )
                self._db.add(chunk)

            # 9. Finalize
            doc.chunk_count = len(chunks)
            doc.status = "ready"
            await self._db.flush()
            # Reload server-generated columns (created_at/updated_at) inside the
            # async context so response serialization doesn't trigger a lazy
            # load outside a greenlet (MissingGreenlet -> 500).
            await self._db.refresh(doc)

        except Exception:
            doc.status = "failed"
            await self._db.flush()
            raise

        return doc

    async def delete(self, doc_id: UUID) -> bool:
        """Soft-delete the Document and remove the encrypted file."""
        doc = await self._doc_repo.get_active(doc_id)
        if doc is None:
            return False

        extension = Path(doc.file_name or "").suffix.lower() or ".bin"
        _delete_stored_file(doc_id, extension)
        await self._chunk_repo.delete_by_document(doc_id)
        await self._doc_repo.soft_delete(doc_id)
        return True

    async def reembed(self, doc_id: UUID) -> Document:
        """
        Re-run chunking + embedding for an existing document.
        Decrypts the stored file, wipes old chunks, re-chunks and re-embeds.
        Raises LookupError if document not found.
        Raises FileNotFoundError if the encrypted file is missing.
        """
        doc = await self._doc_repo.get_active(doc_id)
        if doc is None:
            raise LookupError(f"Document {doc_id} not found")

        extension = Path(doc.file_name or "").suffix.lower() or ".bin"
        raw_bytes = _decrypt_file(doc_id, extension)

        # Wipe old chunks
        await self._chunk_repo.delete_by_document(doc_id)
        doc.status = "processing"
        doc.chunk_count = 0
        await self._db.flush()

        try:
            mime_type = doc.mime_type or "text/plain"
            text = _extract_text(raw_bytes, mime_type)
            text = re.sub(r"\s{3,}", "\n\n", text).strip()

            chunks = _chunk_text(text)
            if not chunks:
                chunks = [text] if text else ["(no content)"]

            # Embed all chunks in one batch call (1 API request per document)
            all_embeddings = await self._emb.embed_batch(chunks)
            actual_dim = len(all_embeddings[0])
            await _ensure_vector_dimension(self._db, actual_dim)
            emb_model = getattr(self._emb, "_model", "nomic-embed-text")

            for idx, (chunk_text, embedding) in enumerate(zip(chunks, all_embeddings)):
                chunk = DocumentChunk(
                    document_id=doc.id,
                    chunk_index=idx,
                    content=chunk_text,
                    content_hash=hashlib.sha256(chunk_text.encode()).hexdigest(),
                    embedding=embedding,
                    embedding_model=emb_model,
                    embedding_dimension=actual_dim,
                    page_number=None,
                )
                self._db.add(chunk)

            doc.chunk_count = len(chunks)
            doc.status = "ready"
            await self._db.flush()

        except Exception:
            doc.status = "failed"
            await self._db.flush()
            raise

        return doc

    async def replace_content(
        self, doc_id: UUID, filename: str, raw_bytes: bytes
    ) -> Document:
        """
        Replace an existing document's stored content with new bytes and
        re-run chunking + embedding. Used when applying a detected source
        update (e.g. a Google Drive file changed). The document id, title and
        category are preserved; the stored file, checksum, size, mime type and
        chunks are refreshed.

        Raises LookupError if the document is not found.
        """
        doc = await self._doc_repo.get_active(doc_id)
        if doc is None:
            raise LookupError(f"Document {doc_id} not found")

        mime_type, _ = mimetypes.guess_type(filename)
        if mime_type not in SUPPORTED_MIME_TYPES:
            ext = Path(filename).suffix.lower()
            mime_type = "text/plain" if ext in (".txt", ".md", ".csv") else (mime_type or "text/plain")

        # Remove the previous encrypted file (extension may change).
        old_ext = Path(doc.file_name or "").suffix.lower() or ".bin"
        try:
            _delete_stored_file(doc_id, old_ext)
        except Exception:
            pass

        # Wipe old chunks.
        await self._chunk_repo.delete_by_document(doc_id)
        doc.status = "processing"
        doc.chunk_count = 0
        doc.file_name = filename
        doc.mime_type = mime_type
        doc.file_size_bytes = len(raw_bytes)
        doc.checksum_sha256 = hashlib.sha256(raw_bytes).hexdigest()
        await self._db.flush()

        try:
            extension = Path(filename).suffix.lower() or ".bin"
            stored_path = _encrypt_and_store(doc_id, raw_bytes, extension)
            doc.source_uri = str(stored_path.relative_to(_get_storage_root()))

            text = _extract_text(raw_bytes, mime_type)
            text = re.sub(r"\s{3,}", "\n\n", text).strip()

            chunks = _chunk_text(text)
            if not chunks:
                chunks = [text] if text else ["(no content)"]

            all_embeddings = await self._emb.embed_batch(chunks)
            actual_dim = len(all_embeddings[0])
            await _ensure_vector_dimension(self._db, actual_dim)
            emb_model = getattr(self._emb, "_model", "nomic-embed-text")

            for idx, (chunk_text, embedding) in enumerate(zip(chunks, all_embeddings)):
                chunk = DocumentChunk(
                    document_id=doc.id,
                    chunk_index=idx,
                    content=chunk_text,
                    content_hash=hashlib.sha256(chunk_text.encode()).hexdigest(),
                    embedding=embedding,
                    embedding_model=emb_model,
                    embedding_dimension=actual_dim,
                    page_number=None,
                )
                self._db.add(chunk)

            doc.chunk_count = len(chunks)
            doc.status = "ready"
            await self._db.flush()
            await self._db.refresh(doc)

        except Exception:
            doc.status = "failed"
            await self._db.flush()
            raise

        return doc


# ── FastAPI dependency ────────────────────────────────────────────────────────

def get_ingestion_service(
    db: AsyncSession,
    embedding_svc: EmbeddingService,
) -> DocumentIngestionService:
    return DocumentIngestionService(db=db, embedding_svc=embedding_svc)
