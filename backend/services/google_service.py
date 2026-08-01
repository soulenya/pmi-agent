"""
Google Workspace integration service.

Handles OAuth token lifecycle and provides read methods for
Gmail, Drive, Calendar, and Contacts plus write methods
(used only after explicit user approval via the proposals API).
"""
from __future__ import annotations

import base64
import email.mime.text
import logging
import threading
import time
from datetime import datetime, timezone, timedelta
from pathlib import Path
from typing import Any

SCOPES = [
    "https://www.googleapis.com/auth/gmail.readonly",
    "https://www.googleapis.com/auth/gmail.send",
    "https://www.googleapis.com/auth/gmail.modify",
    "https://www.googleapis.com/auth/drive.readonly",
    "https://www.googleapis.com/auth/drive.file",
    # Editing a file the user already owns needs the broad Drive scope:
    # drive.file only ever covers files this app itself created. Which files
    # Gerry may actually write to is gated per file in drive_edit_grants.
    "https://www.googleapis.com/auth/drive",
    "https://www.googleapis.com/auth/calendar.readonly",
    "https://www.googleapis.com/auth/calendar.events",
    "https://www.googleapis.com/auth/contacts.readonly",
    "https://www.googleapis.com/auth/spreadsheets.readonly",
    "https://www.googleapis.com/auth/spreadsheets",
    "https://www.googleapis.com/auth/documents.readonly",
    "https://www.googleapis.com/auth/documents",
    # Reading and editing decks. Like Docs, which deck Gerry may write to is
    # gated per file in drive_edit_grants rather than by the scope.
    "https://www.googleapis.com/auth/presentations",
    "https://www.googleapis.com/auth/tasks.readonly",
    "https://www.googleapis.com/auth/userinfo.email",
    "openid",
]

_BACKEND = Path(__file__).parent.parent
CREDS_FILE = _BACKEND / "google_credentials.json"
TOKEN_FILE  = _BACKEND / "google_token.json"

logger = logging.getLogger(__name__)

_auth_status: str = "disconnected"
_auth_lock = threading.Lock()


# ── credential helpers ────────────────────────────────────────────────────

def _log_refresh_failure(exc: Exception) -> None:
    """Token refresh failures were silently swallowed, which made the field
    problem ('Workspace disconnects between sessions') undiagnosable. Leave a
    breadcrumb the feedback diagnostics bundle picks up."""
    try:
        log_dir = _BACKEND / "logs"
        log_dir.mkdir(exist_ok=True)
        with open(log_dir / "google_refresh.log", "a", encoding="utf-8") as f:
            f.write(
                f"{datetime.now(timezone.utc).isoformat()} REFRESH FAILED: "
                f"{type(exc).__name__}: {exc}\n"
            )
    except Exception:
        pass


def get_credentials():
    """Return valid Credentials or None."""
    from google.oauth2.credentials import Credentials
    from google.auth.transport.requests import Request

    if not TOKEN_FILE.exists():
        return None
    # Load with the token's OWN granted scopes (scopes=None). Forcing the
    # current SCOPES list here made EVERY refresh fail with "Not all requested
    # scopes were granted" whenever the stored grant was narrower than the
    # app's list — granular consent checkboxes, or tokens issued by an older
    # build before a scope was added. That was the field cause of "Workspace
    # disconnects between sessions": access tokens live one hour, so the
    # session worked, then the next launch's refresh failed silently. A
    # refresh can never widen a grant anyway; a genuinely missing scope now
    # surfaces as a 403 on the specific call, fixed by reconnecting.
    creds = Credentials.from_authorized_user_file(str(TOKEN_FILE))
    if creds and creds.expired and creds.refresh_token:
        try:
            creds.refresh(Request())
            TOKEN_FILE.write_text(creds.to_json())
        except Exception as exc:  # noqa: BLE001 — log, then report disconnected
            _log_refresh_failure(exc)
            return None
    return creds if (creds and creds.valid) else None


def granted_scopes() -> set[str]:
    """Scopes the stored token actually carries (empty when disconnected).

    A token issued before a scope was added keeps its narrower grant forever —
    refreshing never widens it — so callers that need write access must check
    this rather than assume SCOPES was granted.
    """
    import json

    if not TOKEN_FILE.exists():
        return set()
    try:
        return set(json.loads(TOKEN_FILE.read_text()).get("scopes") or [])
    except Exception:  # noqa: BLE001 — a malformed token is simply no scopes
        return set()


def get_status() -> dict:
    creds = get_credentials()
    if creds:
        import json
        token_data = json.loads(TOKEN_FILE.read_text())
        return {
            "connected": True,
            "status": "connected",
            "email": token_data.get("id_token", {}) if isinstance(token_data.get("id_token"), dict) else "",
        }
    with _auth_lock:
        status = _auth_status
    result: dict = {"connected": False, "status": status}
    if status.startswith("error:"):
        result["error"] = status[6:]
    return result


def start_auth_flow() -> None:
    """Kick off InstalledAppFlow in a background thread — opens browser automatically."""
    global _auth_status
    with _auth_lock:
        if _auth_status == "pending":
            return
        _auth_status = "pending"

    _LOG_FILE = _BACKEND / "logs" / "google_auth.log"

    def _log(msg: str) -> None:
        try:
            _LOG_FILE.parent.mkdir(exist_ok=True)
            import datetime as _dt
            with open(_LOG_FILE, "a", encoding="utf-8") as f:
                f.write(f"{_dt.datetime.now()}: {msg}\n")
        except Exception:
            pass

    def _open_url(url: str, *a, **kw) -> bool:
        """Open a URL from a no-console process using platform-specific fallbacks."""
        _log(f"Opening URL (first 80 chars): {url[:80]}")
        import sys as _sys
        if _sys.platform == "darwin":
            # macOS: hand the URL to the default browser via `open`.
            try:
                import subprocess as _sp
                _sp.Popen(["open", url])
                _log("Browser opened via open")
                return True
            except Exception as e:
                _log(f"open failed: {e}")
            try:
                import webbrowser as _wb
                if _wb.open(url):
                    _log("Browser opened via webbrowser")
                    return True
            except Exception as e:
                _log(f"webbrowser failed: {e}")
            _log("All browser-open attempts failed")
            return False
        # 1. os.startfile — passes URL to Windows shell (default browser handler)
        try:
            import os as _os2
            _os2.startfile(url)
            _log("Browser opened via os.startfile")
            return True
        except Exception as e:
            _log(f"os.startfile failed: {e}")
        # 2. cmd /c start — explicit shell command, always works on Windows
        try:
            import subprocess as _sp
            _sp.Popen(
                ["cmd", "/c", "start", "", url],
                creationflags=_sp.CREATE_NO_WINDOW,
                shell=False,
            )
            _log("Browser opened via cmd /c start")
            return True
        except Exception as e:
            _log(f"cmd /c start failed: {e}")
        # 3. rundll32 — last resort
        try:
            import subprocess as _sp
            _sp.Popen(
                ["rundll32", "url.dll,FileProtocolHandler", url],
                creationflags=_sp.CREATE_NO_WINDOW,
            )
            _log("Browser opened via rundll32")
            return True
        except Exception as e:
            _log(f"rundll32 failed: {e}")
        _log("All browser-open attempts failed")
        return False

    def _run() -> None:
        global _auth_status
        import traceback as _tb
        try:
            _log("Auth thread started")
            import webbrowser
            from google_auth_oauthlib.flow import InstalledAppFlow

            # Patch all webbrowser open variants so run_local_server uses our
            # reliable opener regardless of which method it calls internally.
            webbrowser.open = _open_url
            webbrowser.open_new = _open_url
            webbrowser.open_new_tab = _open_url
            _log("webbrowser patched")

            if not CREDS_FILE.exists():
                raise FileNotFoundError(f"google_credentials.json not found at {CREDS_FILE}")

            flow = InstalledAppFlow.from_client_secrets_file(str(CREDS_FILE), SCOPES)
            _log("Flow created, calling run_local_server...")
            # prompt='consent' forces Google to show ALL scopes every time,
            # bypassing the server-side consent cache from previous partial grants.
            creds = flow.run_local_server(port=0, open_browser=True, prompt="consent")
            _log("run_local_server returned — writing token")
            TOKEN_FILE.write_text(creds.to_json())
            with _auth_lock:
                _auth_status = "connected"
            _log("Auth complete — status = connected")
        except Exception as exc:
            err = str(exc)
            _log(f"EXCEPTION: {err}\n{_tb.format_exc()}")
            with _auth_lock:
                _auth_status = f"error:{err}"

    threading.Thread(target=_run, daemon=True).start()


def revoke() -> None:
    global _auth_status
    if TOKEN_FILE.exists():
        TOKEN_FILE.unlink()
    with _auth_lock:
        _auth_status = "disconnected"


def _require_creds():
    creds = get_credentials()
    if not creds:
        raise RuntimeError("Google account not connected. Please authenticate first.")
    return creds


def _build(service: str, version: str):
    from googleapiclient.discovery import build
    return build(service, version, credentials=_require_creds())


# ── Gmail ─────────────────────────────────────────────────────────────────

def gmail_search(query: str, max_results: int = 10) -> list[dict]:
    svc = _build("gmail", "v1")
    resp = svc.users().messages().list(userId="me", q=query, maxResults=max_results).execute()
    msgs = resp.get("messages", [])
    out = []
    for m in msgs:
        detail = svc.users().messages().get(
            userId="me", id=m["id"], format="metadata",
            metadataHeaders=["From", "To", "Subject", "Date"],
        ).execute()
        headers = {h["name"]: h["value"] for h in detail.get("payload", {}).get("headers", [])}
        out.append({
            "id": m["id"],
            "thread_id": detail.get("threadId", "") or m.get("threadId", ""),
            "from": headers.get("From", ""),
            "to": headers.get("To", ""),
            "subject": headers.get("Subject", ""),
            "date": headers.get("Date", ""),
            "snippet": detail.get("snippet", ""),
        })
    return out


def gmail_get_message(message_id: str) -> dict:
    svc = _build("gmail", "v1")
    m = svc.users().messages().get(userId="me", id=message_id, format="full").execute()
    headers = {h["name"]: h["value"] for h in m.get("payload", {}).get("headers", [])}
    return {
        "id": message_id,
        "from": headers.get("From", ""),
        "to": headers.get("To", ""),
        "subject": headers.get("Subject", ""),
        "date": headers.get("Date", ""),
        "body": _extract_body(m.get("payload", {})),
    }


def gmail_list_drafts(max_results: int = 20) -> list[dict]:
    """List Gmail DRAFTS (unsent) — closes the draft→send blind spot."""
    svc = _build("gmail", "v1")
    resp = svc.users().drafts().list(userId="me", maxResults=max_results).execute()
    out = []
    for d in resp.get("drafts", []):
        msg_id = (d.get("message") or {}).get("id", "")
        item = {"draft_id": d.get("id", ""), "message_id": msg_id,
                "to": "", "subject": "", "date": "", "snippet": ""}
        if msg_id:
            detail = svc.users().messages().get(
                userId="me", id=msg_id, format="metadata",
                metadataHeaders=["To", "Subject", "Date"],
            ).execute()
            headers = {h["name"]: h["value"] for h in detail.get("payload", {}).get("headers", [])}
            item.update(
                to=headers.get("To", ""),
                subject=headers.get("Subject", ""),
                date=headers.get("Date", ""),
                snippet=detail.get("snippet", ""),
            )
        out.append(item)
    return out


def gmail_get_draft(draft_id: str) -> dict:
    """Read one Gmail draft's full content (unsent)."""
    svc = _build("gmail", "v1")
    d = svc.users().drafts().get(userId="me", id=draft_id, format="full").execute()
    m = d.get("message", {}) or {}
    headers = {h["name"]: h["value"] for h in m.get("payload", {}).get("headers", [])}
    return {
        "draft_id": d.get("id", draft_id),
        "message_id": m.get("id", ""),
        "to": headers.get("To", ""),
        "cc": headers.get("Cc", ""),
        "subject": headers.get("Subject", ""),
        "date": headers.get("Date", ""),
        "body": _extract_body(m.get("payload", {})),
    }


def _extract_body(payload: dict) -> str:
    mime = payload.get("mimeType", "")
    if mime == "text/plain":
        data = payload.get("body", {}).get("data", "")
        return base64.urlsafe_b64decode(data + "==").decode("utf-8", errors="ignore")
    if mime.startswith("multipart/"):
        for part in payload.get("parts", []):
            text = _extract_body(part)
            if text:
                return text
    return payload.get("snippet", "")


def _extract_bodies(payload: dict) -> tuple[str, str]:
    """Walk a message payload and return ``(plain_text, html)``.

    Either may be empty. HTML is the raw email HTML (rendered safely in a
    sandboxed iframe on the client).
    """
    text = ""
    html = ""

    def _walk(part: dict) -> None:
        nonlocal text, html
        mime = part.get("mimeType", "")
        data = (part.get("body", {}) or {}).get("data", "")
        if mime == "text/plain" and data and not text:
            text = base64.urlsafe_b64decode(data + "==").decode("utf-8", errors="ignore")
        elif mime == "text/html" and data and not html:
            html = base64.urlsafe_b64decode(data + "==").decode("utf-8", errors="ignore")
        for child in part.get("parts", []) or []:
            _walk(child)

    _walk(payload)
    return text, html


def gmail_send(
    to: str,
    subject: str,
    body: str,
    thread_id: str | None = None,
    reply_to_message_id: str | None = None,
    cc: str | None = None,
    bcc: str | None = None,
    attachments: list[dict] | None = None,
) -> dict:
    """Send a Gmail message.

    When ``reply_to_message_id`` is given, the original message's RFC
    ``Message-ID`` / ``References`` headers are looked up and set as
    ``In-Reply-To`` / ``References`` so the reply threads correctly. When
    ``thread_id`` is given the message is attached to that Gmail thread.
    Threading is best-effort: the email still sends if header lookup fails.

    ``cc`` / ``bcc`` are optional comma-separated address strings. Each item of
    ``attachments`` is ``{filename, mime_type, data (bytes)}``; when any are
    present (or cc/bcc are set) a multipart message is built.
    """
    svc = _build("gmail", "v1")
    attachments = attachments or []
    if attachments:
        import email.mime.multipart
        import email.mime.base
        import email.mime.text as _mt
        from email import encoders as _encoders

        msg: Any = email.mime.multipart.MIMEMultipart()
        msg.attach(_mt.MIMEText(body))
        for att in attachments:
            raw_bytes = att.get("data") or b""
            mime_type = att.get("mime_type") or "application/octet-stream"
            maintype, _, subtype = mime_type.partition("/")
            part = email.mime.base.MIMEBase(maintype or "application", subtype or "octet-stream")
            part.set_payload(raw_bytes)
            _encoders.encode_base64(part)
            part.add_header(
                "Content-Disposition",
                "attachment",
                filename=att.get("filename") or "attachment",
            )
            msg.attach(part)
    else:
        import email.mime.text as _mt

        msg = _mt.MIMEText(body)
    msg["to"] = to
    msg["subject"] = subject
    if cc:
        msg["cc"] = cc
    if bcc:
        msg["bcc"] = bcc
    if reply_to_message_id:
        try:
            orig = (
                svc.users()
                .messages()
                .get(
                    userId="me",
                    id=reply_to_message_id,
                    format="metadata",
                    metadataHeaders=["Message-ID", "References"],
                )
                .execute()
            )
            headers = {
                h["name"].lower(): h["value"]
                for h in orig.get("payload", {}).get("headers", [])
            }
            orig_msg_id = headers.get("message-id", "")
            orig_refs = headers.get("references", "")
            if orig_msg_id:
                msg["In-Reply-To"] = orig_msg_id
                msg["References"] = (orig_refs + " " + orig_msg_id).strip()
        except Exception:
            pass  # threading is best-effort
    raw = base64.urlsafe_b64encode(msg.as_bytes()).decode()
    send_body: dict = {"raw": raw}
    if thread_id:
        send_body["threadId"] = thread_id
    result = svc.users().messages().send(userId="me", body=send_body).execute()
    return {"message_id": result["id"], "status": "sent"}


def _list_attachments(payload: dict) -> list[dict]:
    """List file + inline-image attachments in a message payload (no bytes).

    Each item: ``{filename, mime_type, attachment_id, size, content_id,
    inline}``. Inline images (referenced from HTML via ``cid:``) are included
    so the client can resolve them; ``content_id`` is stripped of <> brackets.
    """
    out: list[dict] = []

    def _walk(part: dict) -> None:
        body = part.get("body", {}) or {}
        att_id = body.get("attachmentId")
        if not att_id:
            for child in part.get("parts", []) or []:
                _walk(child)
            return
        headers = {h["name"].lower(): h["value"] for h in part.get("headers", []) or []}
        content_id = (headers.get("content-id") or "").strip().strip("<>")
        disposition = headers.get("content-disposition", "")
        mime = part.get("mimeType", "")
        filename = part.get("filename") or ""
        inline = bool(content_id) or disposition.lower().startswith("inline")
        if not filename:
            # Inline image without a filename — synthesize one from the MIME type.
            ext = mime.split("/")[-1] if "/" in mime else "bin"
            filename = f"image.{ext}" if mime.startswith("image/") else f"attachment.{ext}"
        out.append({
            "filename": filename,
            "mime_type": mime,
            "attachment_id": att_id,
            "size": body.get("size", 0),
            "content_id": content_id,
            "inline": inline,
        })
        for child in part.get("parts", []) or []:
            _walk(child)

    _walk(payload)
    return out


def gmail_list_threads(query: str = "in:inbox", max_results: int = 25) -> list[dict]:
    """List threads as lightweight summaries for the Inbox view.

    Each item: ``{thread_id, subject, from, date, snippet, message_count,
    unread}``. ``from``/``date`` come from the latest message in the thread.
    """
    svc = _build("gmail", "v1")
    resp = svc.users().threads().list(
        userId="me", q=query, maxResults=max_results
    ).execute()
    out: list[dict] = []
    for t in resp.get("threads", []):
        detail = svc.users().threads().get(
            userId="me", id=t["id"], format="metadata",
            metadataHeaders=["From", "Subject", "Date"],
        ).execute()
        msgs = detail.get("messages", [])
        first = msgs[0] if msgs else {}
        last = msgs[-1] if msgs else {}
        first_h = {h["name"]: h["value"] for h in first.get("payload", {}).get("headers", [])}
        last_h = {h["name"]: h["value"] for h in last.get("payload", {}).get("headers", [])}
        labels: set[str] = set()
        for m in msgs:
            labels.update(m.get("labelIds", []) or [])
        out.append({
            "thread_id": t["id"],
            "subject": first_h.get("Subject", "") or last_h.get("Subject", ""),
            "from": last_h.get("From", ""),
            "date": last_h.get("Date", ""),
            "snippet": t.get("snippet", ""),
            "message_count": len(msgs),
            "unread": "UNREAD" in labels,
        })
    return out


def gmail_get_thread(thread_id: str) -> dict:
    """Return every message in a Gmail thread with bodies + attachment metadata.

    Shape: ``{thread_id, subject, messages: [{id, from, to, subject, date,
    body, attachments: [{filename, mime_type, attachment_id, size}]}]}``.
    Attachment *bytes* are not downloaded here — use ``gmail_get_attachments``
    at import time. Used by the "Add email thread to Knowledge Base" feature.
    """
    svc = _build("gmail", "v1")
    thread = svc.users().threads().get(userId="me", id=thread_id, format="full").execute()
    messages: list[dict] = []
    subject = ""
    for m in thread.get("messages", []):
        payload = m.get("payload", {})
        headers = {h["name"]: h["value"] for h in payload.get("headers", [])}
        if not subject:
            subject = headers.get("Subject", "")
        text, html = _extract_bodies(payload)
        messages.append({
            "id": m.get("id", ""),
            "from": headers.get("From", ""),
            "to": headers.get("To", ""),
            "cc": headers.get("Cc", ""),
            "subject": headers.get("Subject", ""),
            "date": headers.get("Date", ""),
            "body": text or _extract_body(payload),
            "body_html": html,
            "attachments": _list_attachments(payload),
        })
    # Every address this account owns, so the UI can exclude them from
    # Reply-all. The profile address alone is not enough: mail addressed to a
    # send-as alias (morganjkeane@precisianmedical.com vs the primary
    # morganjkeane@pmi-llc.com) was being Cc'd back to the user.
    own = gmail_own_addresses()
    me = own[0] if own else ""
    return {
        "thread_id": thread_id,
        "subject": subject,
        "me": me,
        "me_addresses": own,
        "messages": messages,
    }


def gmail_trash_thread(thread_id: str) -> dict:
    """Move an entire thread to the Gmail Trash (recoverable for 30 days).

    Requires the ``gmail.modify`` scope. Returns ``{thread_id, status}``.
    """
    svc = _build("gmail", "v1")
    svc.users().threads().trash(userId="me", id=thread_id).execute()
    return {"thread_id": thread_id, "status": "trashed"}


def gmail_profile_email() -> str:
    """Return the connected Gmail account's email address ('' on failure)."""
    try:
        svc = _build("gmail", "v1")
        return svc.users().getProfile(userId="me").execute().get("emailAddress", "")
    except Exception:
        return ""


_own_addresses_cache: tuple[float, list[str]] | None = None


def gmail_own_addresses() -> list[str]:
    """Every address the connected account sends as — profile email first.

    ``settings.sendAs.list`` is covered by the existing gmail.readonly scope,
    so no re-consent is needed. Cached for an hour; aliases rarely change.
    """
    global _own_addresses_cache
    now = time.time()
    if _own_addresses_cache and now - _own_addresses_cache[0] < 3600:
        return list(_own_addresses_cache[1])

    addresses: list[str] = []
    primary = gmail_profile_email()
    if primary:
        addresses.append(primary.lower())
    try:
        svc = _build("gmail", "v1")
        resp = svc.users().settings().sendAs().list(userId="me").execute()
        for sa in resp.get("sendAs", []):
            addr = (sa.get("sendAsEmail") or "").strip().lower()
            if addr and addr not in addresses:
                addresses.append(addr)
    except Exception:
        logger.exception("Could not list send-as aliases")

    if addresses:
        _own_addresses_cache = (now, addresses)
    return list(addresses)


def gmail_get_signature() -> str:
    """Return the HTML signature of the primary send-as address ('' if none).

    Reads ``users.settings.sendAs`` — covered by the existing gmail.readonly
    scope, so no re-consent is required. The ``sendAs.list`` response does not
    reliably populate the ``signature`` field, so we fall back to fetching each
    alias individually (primary first) with ``sendAs.get``.
    """
    svc = _build("gmail", "v1")
    resp = svc.users().settings().sendAs().list(userId="me").execute()
    send_as = resp.get("sendAs", [])

    # Primary alias first, then the rest.
    ordered = sorted(send_as, key=lambda sa: not sa.get("isPrimary"))

    # 1) Use any signature already present in the list response.
    for sa in ordered:
        if sa.get("signature"):
            return sa["signature"]

    # 2) list() often omits the signature — fetch each alias to retrieve it.
    for sa in ordered:
        addr = sa.get("sendAsEmail")
        if not addr:
            continue
        try:
            full = (
                svc.users()
                .settings()
                .sendAs()
                .get(userId="me", sendAsEmail=addr)
                .execute()
            )
        except Exception:
            continue
        if full.get("signature"):
            return full["signature"]

    return ""


def gmail_get_attachment(message_id: str, attachment_id: str) -> bytes:
    """Download the raw bytes of a single Gmail attachment."""
    svc = _build("gmail", "v1")
    att = (
        svc.users()
        .messages()
        .attachments()
        .get(userId="me", messageId=message_id, id=attachment_id)
        .execute()
    )
    data = att.get("data", "")
    return base64.urlsafe_b64decode(data + "==") if data else b""


def gmail_get_attachments(message_id: str) -> list[dict]:
    """Return downloadable file attachments for a Gmail message.

    Each item: ``{filename, mime_type, attachment_id, size, data (bytes)}``.
    Inline parts without a filename are skipped. Used by the daily assistant
    scan to auto-import meeting-summary attachments into the Knowledge Base.
    """
    svc = _build("gmail", "v1")
    msg = svc.users().messages().get(userId="me", id=message_id, format="full").execute()
    out: list[dict] = []

    def _walk(part: dict) -> None:
        filename = part.get("filename") or ""
        body = part.get("body", {}) or {}
        att_id = body.get("attachmentId")
        if filename and att_id:
            att = (
                svc.users()
                .messages()
                .attachments()
                .get(userId="me", messageId=message_id, id=att_id)
                .execute()
            )
            data = att.get("data", "")
            raw = base64.urlsafe_b64decode(data + "==") if data else b""
            out.append({
                "filename": filename,
                "mime_type": part.get("mimeType", ""),
                "attachment_id": att_id,
                "size": body.get("size", 0),
                "data": raw,
            })
        for child in part.get("parts", []) or []:
            _walk(child)

    _walk(msg.get("payload", {}))
    return out



# ── Drive ─────────────────────────────────────────────────────────────────

def drive_search(query: str, max_results: int = 10) -> list[dict]:
    svc = _build("drive", "v3")
    resp = svc.files().list(
        q=f"fullText contains '{query}' and trashed=false",
        pageSize=max_results,
        fields="files(id,name,mimeType,modifiedTime,webViewLink,owners)",
        includeItemsFromAllDrives=True,
        supportsAllDrives=True,
        corpora="allDrives",
    ).execute()
    return [
        {
            "id": f["id"],
            "name": f["name"],
            "type": f.get("mimeType", ""),
            "modified": f.get("modifiedTime", ""),
            "url": f.get("webViewLink", ""),
            "owner": (f.get("owners") or [{}])[0].get("displayName", ""),
        }
        for f in resp.get("files", [])
    ]


def drive_recent_files(max_results: int = 8) -> list[dict]:
    """The user's most recently modified Drive files (docs first, newest first)."""
    svc = _build("drive", "v3")
    resp = svc.files().list(
        q="trashed=false and mimeType != 'application/vnd.google-apps.folder'",
        orderBy="modifiedTime desc",
        pageSize=max_results,
        fields="files(id,name,mimeType,modifiedTime,webViewLink,owners)",
        includeItemsFromAllDrives=True,
        supportsAllDrives=True,
        corpora="allDrives",
    ).execute()
    return [
        {
            "id": f["id"],
            "name": f["name"],
            "type": f.get("mimeType", ""),
            "modified": f.get("modifiedTime", ""),
            "url": f.get("webViewLink", ""),
            "owner": (f.get("owners") or [{}])[0].get("displayName", ""),
        }
        for f in resp.get("files", [])
    ]


def drive_upload_file(
    local_path: str,
    name: str | None = None,
    folder_id: str | None = None,
) -> dict:
    """Upload a local file to the user's Google Drive and return its metadata.

    Requires the ``drive.file`` scope (granted on connect). ``name`` overrides the
    Drive filename; ``folder_id`` places the file inside a folder. Returns
    ``{id, name, url}`` where ``url`` is the shareable webViewLink.
    """
    import mimetypes
    from googleapiclient.http import MediaFileUpload

    p = Path(local_path)
    if not p.is_file():
        raise RuntimeError(f"Local file not found: {local_path}")

    drive_name = name or p.name
    mime_type = mimetypes.guess_type(p.name)[0] or "application/octet-stream"

    metadata: dict[str, Any] = {"name": drive_name}
    if folder_id:
        metadata["parents"] = [folder_id]

    svc = _build("drive", "v3")
    media = MediaFileUpload(str(p), mimetype=mime_type, resumable=False)
    created = svc.files().create(
        body=metadata,
        media_body=media,
        fields="id,name,webViewLink",
        supportsAllDrives=True,
    ).execute()
    return {
        "id": created.get("id", ""),
        "name": created.get("name", drive_name),
        "url": created.get("webViewLink", ""),
    }


def drive_upload_bytes(
    data: bytes,
    name: str,
    mime_type: str | None = None,
    folder_id: str | None = None,
) -> dict:
    """Upload in-memory bytes to the user's Google Drive and return its metadata.

    ``folder_id`` places the file inside a folder (My Drive root when omitted).
    Returns ``{id, name, url}`` where ``url`` is the shareable webViewLink.
    """
    import io
    import mimetypes
    from googleapiclient.http import MediaIoBaseUpload

    mt = mime_type or mimetypes.guess_type(name)[0] or "application/octet-stream"
    metadata: dict[str, Any] = {"name": name}
    if folder_id:
        metadata["parents"] = [folder_id]

    svc = _build("drive", "v3")
    media = MediaIoBaseUpload(io.BytesIO(data), mimetype=mt, resumable=False)
    created = svc.files().create(
        body=metadata,
        media_body=media,
        fields="id,name,webViewLink",
        supportsAllDrives=True,
    ).execute()
    return {
        "id": created.get("id", ""),
        "name": created.get("name", name),
        "url": created.get("webViewLink", ""),
    }


# Office / text formats Google Drive can convert into native Workspace docs.
_WORKSPACE_IMPORT_MAP = {
    # → Google Docs
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        "application/vnd.google-apps.document",
    "application/msword": "application/vnd.google-apps.document",
    "application/rtf": "application/vnd.google-apps.document",
    "text/rtf": "application/vnd.google-apps.document",
    "application/vnd.oasis.opendocument.text": "application/vnd.google-apps.document",
    "text/plain": "application/vnd.google-apps.document",
    "text/html": "application/vnd.google-apps.document",
    # → Google Sheets
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
        "application/vnd.google-apps.spreadsheet",
    "application/vnd.ms-excel": "application/vnd.google-apps.spreadsheet",
    "application/vnd.oasis.opendocument.spreadsheet": "application/vnd.google-apps.spreadsheet",
    "text/csv": "application/vnd.google-apps.spreadsheet",
    "text/tab-separated-values": "application/vnd.google-apps.spreadsheet",
    # → Google Slides
    "application/vnd.openxmlformats-officedocument.presentationml.presentation":
        "application/vnd.google-apps.presentation",
    "application/vnd.ms-powerpoint": "application/vnd.google-apps.presentation",
    "application/vnd.oasis.opendocument.presentation": "application/vnd.google-apps.presentation",
}


def drive_import_attachment(
    data: bytes, name: str, source_mime: str | None = None
) -> dict:
    """Upload bytes to Drive so they open in Google Workspace, and return the link.

    Office/text files (docx, xlsx, pptx, csv, txt, …) are converted into the
    matching native Google doc (Docs/Sheets/Slides) so they open in the Workspace
    editor. Non-convertible types (PDF, images, archives, …) are uploaded as-is
    and open in the Drive viewer. Requires the ``drive.file`` scope (granted on
    connect). Returns ``{id, name, url}`` where ``url`` is the webViewLink.
    """
    import io
    import mimetypes
    from googleapiclient.http import MediaIoBaseUpload

    src = (source_mime or mimetypes.guess_type(name)[0] or "application/octet-stream")
    src_key = src.split(";")[0].strip().lower()
    target = _WORKSPACE_IMPORT_MAP.get(src_key)

    metadata: dict[str, Any] = {"name": name}
    if target:
        metadata["mimeType"] = target

    svc = _build("drive", "v3")
    media = MediaIoBaseUpload(io.BytesIO(data), mimetype=src_key, resumable=False)
    created = svc.files().create(
        body=metadata,
        media_body=media,
        fields="id,name,webViewLink",
        supportsAllDrives=True,
    ).execute()
    return {
        "id": created.get("id", ""),
        "name": created.get("name", name),
        "url": created.get("webViewLink", ""),
    }


def drive_ocr_extract_text(data: bytes, name: str, source_mime: str | None = None) -> str:
    """OCR image/PDF bytes by converting them into a TEMPORARY Google Doc.

    The temp doc is app-created (rides the ``drive.file`` scope), its text is
    exported, and it is trashed immediately — nothing is left behind. Returns
    the extracted text ("" when Google couldn't OCR the content).
    """
    import io
    import mimetypes
    from googleapiclient.http import MediaIoBaseUpload

    src = (source_mime or mimetypes.guess_type(name)[0] or "application/octet-stream")
    svc = _build("drive", "v3")
    media = MediaIoBaseUpload(io.BytesIO(data), mimetype=src.split(";")[0].strip(), resumable=False)
    created = svc.files().create(
        body={"name": f"[gerry-ocr-temp] {name}"[:200],
              "mimeType": "application/vnd.google-apps.document"},
        media_body=media,
        fields="id",
    ).execute()
    temp_id = created.get("id", "")
    if not temp_id:
        return ""
    try:
        raw = svc.files().export(fileId=temp_id, mimeType="text/plain").execute()
        return raw.decode("utf-8", errors="ignore") if isinstance(raw, bytes) else str(raw or "")
    finally:
        try:
            svc.files().update(fileId=temp_id, body={"trashed": True}).execute()
        except Exception:  # noqa: BLE001 — a stray temp doc is cosmetic, not fatal
            logger.warning("Failed to trash OCR temp doc %s", temp_id)


def drive_find_file_matches(name: str, max_results: int = 25) -> list[dict]:
    """Find non-trashed Drive files whose name EXACTLY equals ``name``.

    Returns candidates with id, name, size (int|None), type, modified, url —
    newest first. Used to relink locally-uploaded Knowledge Base documents to
    their Drive original so they become update-trackable and shareable.
    """
    svc = _build("drive", "v3")
    safe = name.replace("\\", "\\\\").replace("'", "\\'")
    resp = svc.files().list(
        q=f"name = '{safe}' and trashed=false",
        pageSize=max_results,
        orderBy="modifiedTime desc",
        fields="files(id,name,size,mimeType,modifiedTime,webViewLink)",
        includeItemsFromAllDrives=True,
        supportsAllDrives=True,
        corpora="allDrives",
    ).execute()
    out: list[dict] = []
    for f in resp.get("files", []):
        sz = f.get("size")
        out.append(
            {
                "id": f["id"],
                "name": f["name"],
                "size": int(sz) if sz is not None else None,
                "type": f.get("mimeType", ""),
                "modified": f.get("modifiedTime", ""),
                "url": f.get("webViewLink", ""),
            }
        )
    return out


def drive_search_by_name(name_contains: str, max_results: int = 25) -> list[dict]:
    """Find Drive files whose name contains ``name_contains`` (newest first).

    Used by the daily assistant scan to locate Gemini meeting-notes Docs, which
    Google names like ``"<Meeting> - <date> - Notes by Gemini"``.
    """
    svc = _build("drive", "v3")
    safe = name_contains.replace("\\", "\\\\").replace("'", "\\'")
    resp = svc.files().list(
        q=f"name contains '{safe}' and trashed=false",
        pageSize=max_results,
        orderBy="modifiedTime desc",
        fields="files(id,name,mimeType,modifiedTime,webViewLink,owners)",
        includeItemsFromAllDrives=True,
        supportsAllDrives=True,
        corpora="allDrives",
    ).execute()
    return [
        {
            "id": f["id"],
            "name": f["name"],
            "type": f.get("mimeType", ""),
            "modified": f.get("modifiedTime", ""),
            "url": f.get("webViewLink", ""),
            "owner": (f.get("owners") or [{}])[0].get("displayName", ""),
        }
        for f in resp.get("files", [])
    ]


def drive_get_file_meta(file_id: str) -> dict:
    """Fetch minimal metadata for a Drive item: {id, name, mime_type}."""
    svc = _build("drive", "v3")
    meta = svc.files().get(
        fileId=file_id, fields="id,name,mimeType", supportsAllDrives=True
    ).execute()
    return {"id": meta.get("id", file_id), "name": meta.get("name", ""), "mime_type": meta.get("mimeType", "")}


def drive_list_folder(
    folder_id: str = "root",
    drive_id: str | None = None,
    max_results: int = 100,
) -> list[dict]:
    """List the direct children (files and folders) of a Drive folder, including shared drives.

    When ``drive_id`` is supplied the caller is navigating the ROOT of a shared
    drive.  The Google Drive API requires ``corpora='drive'`` + ``driveId`` in
    that case; the regular ``'{id}' in parents`` query returns nothing.
    """
    svc = _build("drive", "v3")

    # Shared-drive root: use corpora='drive' mode
    if drive_id:
        resp = svc.files().list(
            corpora="drive",
            driveId=drive_id,
            q=f"'{folder_id}' in parents and trashed=false",
            pageSize=max_results,
            orderBy="folder,name",
            fields="files(id,name,mimeType,modifiedTime,webViewLink)",
            includeItemsFromAllDrives=True,
            supportsAllDrives=True,
        ).execute()
    else:
        resp = svc.files().list(
            q=f"'{folder_id}' in parents and trashed=false",
            pageSize=max_results,
            orderBy="folder,name",
            fields="files(id,name,mimeType,modifiedTime,webViewLink)",
            includeItemsFromAllDrives=True,
            supportsAllDrives=True,
        ).execute()
    items = []
    for f in resp.get("files", []):
        mime = f.get("mimeType", "")
        is_folder = mime == "application/vnd.google-apps.folder"
        items.append({
            "id": f["id"],
            "name": f["name"],
            "type": "folder" if is_folder else mime,
            "modified": f.get("modifiedTime", ""),
            "url": f.get("webViewLink", ""),
        })
    return items


def drive_list_shared_drives(max_results: int = 20) -> list[dict]:
    """List all shared/team drives the user has access to."""
    svc = _build("drive", "v3")
    resp = svc.drives().list(
        pageSize=max_results,
        fields="drives(id,name)",
    ).execute()
    return [
        {"id": d["id"], "name": d["name"], "type": "shared_drive"}
        for d in resp.get("drives", [])
    ]


def drive_find_or_create_folder(name: str, parent_id: str | None = None) -> dict:
    """Find a folder by exact name (across all drives) or create it.

    ``parent_id`` scopes creation (a folder id, or a shared-drive id to create
    in that drive's root). Returns ``{id, name, url, created}``.
    """
    svc = _build("drive", "v3")
    safe = name.replace("'", "\\'")
    q = (
        f"name = '{safe}' and mimeType = 'application/vnd.google-apps.folder' "
        "and trashed = false"
    )
    resp = svc.files().list(
        q=q,
        pageSize=5,
        fields="files(id,name,webViewLink)",
        includeItemsFromAllDrives=True,
        supportsAllDrives=True,
        corpora="allDrives",
    ).execute()
    files = resp.get("files", [])
    if files:
        f = files[0]
        return {"id": f["id"], "name": f["name"], "url": f.get("webViewLink", ""), "created": False}

    metadata: dict[str, Any] = {
        "name": name,
        "mimeType": "application/vnd.google-apps.folder",
    }
    if parent_id:
        metadata["parents"] = [parent_id]
    created = svc.files().create(
        body=metadata,
        fields="id,name,webViewLink",
        supportsAllDrives=True,
    ).execute()
    return {
        "id": created.get("id", ""),
        "name": created.get("name", name),
        "url": created.get("webViewLink", ""),
        "created": True,
    }


def drive_update_bytes(file_id: str, data: bytes, mime_type: str | None = None) -> dict:
    """Overwrite an existing Drive file's content in place, keeping its id.

    Returns ``{id, name, url}``. Raises googleapiclient HttpError on 404 so
    callers can fall back to creating a fresh file.
    """
    import io
    from googleapiclient.http import MediaIoBaseUpload

    svc = _build("drive", "v3")
    media = MediaIoBaseUpload(
        io.BytesIO(data),
        mimetype=mime_type or "application/octet-stream",
        resumable=False,
    )
    updated = svc.files().update(
        fileId=file_id,
        media_body=media,
        fields="id,name,webViewLink",
        supportsAllDrives=True,
    ).execute()
    return {
        "id": updated.get("id", file_id),
        "name": updated.get("name", ""),
        "url": updated.get("webViewLink", ""),
    }


def drive_get_content(file_id: str, max_chars: int | None = 10_000) -> dict:
    """Fetch a Drive file's text content.

    ``max_chars`` caps the returned ``content`` (None = full text). The result
    always reports ``total_chars`` (pre-cap length) and ``truncated`` so callers
    can surface an explicit marker instead of silently cutting mid-sentence.
    """
    svc = _build("drive", "v3")
    meta = svc.files().get(
        fileId=file_id,
        fields="id,name,mimeType,webViewLink,modifiedTime",
        supportsAllDrives=True,
    ).execute()
    mime = meta.get("mimeType", "")
    content = ""
    raw_file_bytes: bytes | None = None
    file_extension = ""

    export_map = {
        "application/vnd.google-apps.document":     "text/plain",
        "application/vnd.google-apps.spreadsheet":  "text/csv",
        "application/vnd.google-apps.presentation": "text/plain",
    }
    if mime in export_map:
        raw = svc.files().export(fileId=file_id, mimeType=export_map[mime]).execute()
        content = raw.decode("utf-8", errors="ignore") if isinstance(raw, bytes) else str(raw)
    elif mime.startswith("text/") or mime in ("application/json",):
        raw = svc.files().get_media(fileId=file_id).execute()
        content = raw.decode("utf-8", errors="ignore") if isinstance(raw, bytes) else str(raw)
    elif mime == "application/pdf":
        # Return the raw PDF bytes (for the import path to extract via PyMuPDF,
        # identical to uploads) and also extract text here with PyMuPDF for the
        # agent reader and update-sync. pypdf (used previously) silently returned
        # empty text on many PDFs, so Drive imports failed where the identical
        # uploaded file worked.
        raw_bytes = svc.files().get_media(fileId=file_id).execute()
        if isinstance(raw_bytes, bytes):
            raw_file_bytes = raw_bytes
            file_extension = ".pdf"
            try:
                import fitz  # PyMuPDF
                with fitz.open(stream=raw_bytes, filetype="pdf") as pdf:
                    content = "\n\n".join(
                        t for t in (page.get_text() for page in pdf) if t.strip()
                    )
            except Exception:
                content = ""
    elif mime in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    ):
        # Uploaded Word docs are NOT Google-native, so files().export() returns
        # 403 fileNotExportable. Return the raw bytes for the import path to parse
        # with python-docx (identical to the upload path) and also extract text
        # here for the agent reader and update-sync.
        raw_bytes = svc.files().get_media(fileId=file_id).execute()
        if isinstance(raw_bytes, bytes):
            raw_file_bytes = raw_bytes
            file_extension = ".docx"
            try:
                import io
                import docx
                document = docx.Document(io.BytesIO(raw_bytes))
                parts = [p.text for p in document.paragraphs if p.text.strip()]
                # Include table cell text, common in cover letters/forms.
                for table in document.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                parts.append(cell.text)
                content = "\n".join(parts)
            except Exception:
                content = ""

    total_chars = len(content)
    was_truncated = max_chars is not None and total_chars > max_chars
    if was_truncated:
        content = content[:max_chars]

    return {
        "id": file_id,
        "name": meta.get("name", ""),
        "type": mime,
        "url": meta.get("webViewLink", ""),
        "modified": meta.get("modifiedTime", ""),
        "content": content,
        "truncated": was_truncated,
        "total_chars": total_chars,
        "raw_bytes": raw_file_bytes,
        "extension": file_extension,
    }


def drive_download_bytes(file_id: str) -> dict:
    """Download a Drive file's raw bytes for storage in the regulatory file store.

    Google-native files (Docs/Sheets/Slides) are exported to their Office
    equivalents (.docx/.xlsx/.pptx); everything else is downloaded as-is via
    get_media. Returns ``{name, mime_type, extension, content (bytes),
    modified, url}``. The returned ``name`` already carries the right extension.
    """
    svc = _build("drive", "v3")
    meta = svc.files().get(
        fileId=file_id,
        fields="id,name,mimeType,webViewLink,modifiedTime",
        supportsAllDrives=True,
    ).execute()
    mime = meta.get("mimeType", "")
    name = meta.get("name", "file")

    # Google-native → export to an Office format
    google_export = {
        "application/vnd.google-apps.document": (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ".docx",
        ),
        "application/vnd.google-apps.spreadsheet": (
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            ".xlsx",
        ),
        "application/vnd.google-apps.presentation": (
            "application/vnd.openxmlformats-officedocument.presentationml.presentation",
            ".pptx",
        ),
    }

    if mime in google_export:
        out_mime, ext = google_export[mime]
        raw = svc.files().export(fileId=file_id, mimeType=out_mime).execute()
        if not name.lower().endswith(ext):
            name = f"{name}{ext}"
        content = raw if isinstance(raw, bytes) else str(raw).encode("utf-8")
        result_mime = out_mime
    else:
        raw = svc.files().get_media(fileId=file_id).execute()
        content = raw if isinstance(raw, bytes) else str(raw).encode("utf-8")
        result_mime = mime
        ext = ""
        if "." in name:
            ext = name[name.rfind("."):]

    return {
        "name": name,
        "mime_type": result_mime,
        "extension": ext.lower(),
        "content": content,
        "modified": meta.get("modifiedTime", ""),
        "url": meta.get("webViewLink", ""),
    }


def drive_list_comments(file_id: str, max_results: int = 50) -> list[dict]:
    """List comment threads on a Drive file (any type) — author, anchor text,
    content, resolved state, and replies. Deleted comments are excluded."""
    svc = _build("drive", "v3")
    out: list[dict] = []
    page_token = None
    while len(out) < max_results:
        resp = svc.comments().list(
            fileId=file_id,
            fields=(
                "nextPageToken,comments(author(displayName),content,"
                "quotedFileContent(value),resolved,createdTime,"
                "replies(author(displayName),content,createdTime))"
            ),
            pageSize=min(100, max_results),
            pageToken=page_token,
            includeDeleted=False,
        ).execute()
        for c in resp.get("comments", []):
            out.append(
                {
                    "author": (c.get("author") or {}).get("displayName", ""),
                    "content": c.get("content", ""),
                    "anchor": (c.get("quotedFileContent") or {}).get("value", ""),
                    "resolved": bool(c.get("resolved")),
                    "created": (c.get("createdTime") or "")[:10],
                    "replies": [
                        {
                            "author": (r.get("author") or {}).get("displayName", ""),
                            "content": r.get("content", ""),
                            "created": (r.get("createdTime") or "")[:10],
                        }
                        for r in c.get("replies", [])
                    ],
                }
            )
        page_token = resp.get("nextPageToken")
        if not page_token:
            break
    return out[:max_results]


def docs_get_suggestions(file_id: str) -> dict:
    """Suggested edits on a NATIVE Google Doc, rendered per paragraph with
    inline markers: {++suggested insertion++} / {--suggested deletion--}.

    Only paragraphs containing suggestions are returned. Raises HttpError for
    non-Docs files (400/404) — callers translate that honestly. Note: the Docs
    API does not expose WHO made a suggestion, only its content.
    """
    svc = _build("docs", "v1")
    doc = svc.documents().get(
        documentId=file_id, suggestionsViewMode="SUGGESTIONS_INLINE"
    ).execute()

    paragraphs: list[str] = []
    for element in (doc.get("body") or {}).get("content", []):
        para = element.get("paragraph")
        if not para:
            continue
        parts: list[str] = []
        has_suggestion = False
        for pe in para.get("elements", []):
            run = pe.get("textRun")
            if not run:
                continue
            text = run.get("content", "")
            if run.get("suggestedInsertionIds"):
                parts.append("{++" + text.rstrip("\n") + "++}")
                has_suggestion = True
            elif run.get("suggestedDeletionIds"):
                parts.append("{--" + text.rstrip("\n") + "--}")
                has_suggestion = True
            else:
                parts.append(text)
        if has_suggestion:
            paragraphs.append("".join(parts).strip())
    return {"title": doc.get("title", ""), "paragraphs": paragraphs}


def _docs_body_end(svc, file_id: str) -> tuple[str, int]:
    """(title, endIndex of the body) for a native Google Doc."""
    doc = svc.documents().get(
        documentId=file_id, fields="title,body(content(endIndex))"
    ).execute()
    content = (doc.get("body") or {}).get("content") or []
    end = content[-1].get("endIndex", 1) if content else 1
    return doc.get("title", ""), int(end)


def docs_append_text(file_id: str, text: str) -> dict:
    """Append text to the end of a native Google Doc. Returns ``{title, chars}``."""
    svc = _build("docs", "v1")
    title, end = _docs_body_end(svc, file_id)
    # The segment's trailing newline is not a legal insertion index.
    svc.documents().batchUpdate(
        documentId=file_id,
        body={"requests": [{"insertText": {"location": {"index": max(1, end - 1)}, "text": text}}]},
    ).execute()
    return {"title": title, "chars": len(text)}


def docs_replace_text(file_id: str, find: str, replace: str, match_case: bool = True) -> dict:
    """Replace every occurrence of ``find`` in a native Google Doc.

    Returns ``{title, occurrences}``; 0 occurrences is not an error.
    """
    svc = _build("docs", "v1")
    resp = svc.documents().batchUpdate(
        documentId=file_id,
        body={
            "requests": [
                {
                    "replaceAllText": {
                        "containsText": {"text": find, "matchCase": match_case},
                        "replaceText": replace,
                    }
                }
            ]
        },
    ).execute()
    replies = resp.get("replies") or [{}]
    changed = int((replies[0].get("replaceAllText") or {}).get("occurrencesChanged", 0) or 0)
    meta = svc.documents().get(documentId=file_id, fields="title").execute()
    return {"title": meta.get("title", ""), "occurrences": changed}


def docs_overwrite_text(file_id: str, text: str) -> dict:
    """Replace a native Google Doc's entire body with ``text``.

    Destructive by design — callers must have a per-file grant, and Drive's own
    version history is the undo path.
    """
    svc = _build("docs", "v1")
    title, end = _docs_body_end(svc, file_id)
    requests: list[dict] = []
    if end > 2:
        requests.append({"deleteContentRange": {"range": {"startIndex": 1, "endIndex": end - 1}}})
    requests.append({"insertText": {"location": {"index": 1}, "text": text}})
    svc.documents().batchUpdate(documentId=file_id, body={"requests": requests}).execute()
    return {"title": title, "chars": len(text)}


def _slides_shape_text(shape_el: dict) -> str:
    """Concatenate the text runs of one page element."""
    parts: list[str] = []
    for el in ((shape_el.get("shape") or {}).get("text") or {}).get("textElements") or []:
        run = el.get("textRun") or {}
        if run.get("content"):
            parts.append(run["content"])
    return "".join(parts).strip()


def slides_read(file_id: str) -> dict:
    """Read a Google Slides deck as plain structure.

    Returns ``{title, slide_count, slides: [{index, object_id, text: [...]}]}``.
    Text is per shape and in reading order, which is what the model needs to
    reason about a deck without seeing it.
    """
    svc = _build("slides", "v1")
    deck = svc.presentations().get(presentationId=file_id).execute()
    slides = []
    for i, page in enumerate(deck.get("slides") or []):
        texts = []
        for el in page.get("pageElements") or []:
            body = _slides_shape_text(el)
            if body:
                texts.append({"object_id": el.get("objectId", ""), "text": body})
        slides.append({
            "index": i + 1,
            "object_id": page.get("objectId", ""),
            "text": texts,
        })
    return {
        "title": deck.get("title", ""),
        "slide_count": len(slides),
        "slides": slides,
    }


def slides_replace_text(file_id: str, find: str, replace: str, match_case: bool = True) -> dict:
    """Replace every occurrence of ``find`` across a deck. Returns occurrences changed."""
    svc = _build("slides", "v1")
    result = svc.presentations().batchUpdate(
        presentationId=file_id,
        body={"requests": [{
            "replaceAllText": {
                "containsText": {"text": find, "matchCase": match_case},
                "replaceText": replace,
            }
        }]},
    ).execute()
    replies = result.get("replies") or [{}]
    changed = int((replies[0].get("replaceAllText") or {}).get("occurrencesChanged", 0) or 0)
    meta = svc.presentations().get(presentationId=file_id, fields="title").execute()
    return {"title": meta.get("title", ""), "occurrences": changed}


_TEXT_STYLE_FIELDS = (
    "backgroundColor", "baselineOffset", "bold", "fontFamily", "fontSize",
    "foregroundColor", "italic", "smallCaps", "strikethrough", "underline",
    "weightedFontFamily",
)
_PARA_STYLE_FIELDS = (
    "alignment", "direction", "indentEnd", "indentFirstLine", "indentStart",
    "lineSpacing", "spaceAbove", "spaceBelow", "spacingMode",
)


def _slides_find_element(deck: dict, object_id: str) -> dict | None:
    """Locate a page element by id, descending into groups."""
    def walk(elements: list[dict]) -> dict | None:
        for el in elements:
            if el.get("objectId") == object_id:
                return el
            children = (el.get("elementGroup") or {}).get("children") or []
            found = walk(children)
            if found:
                return found
        return None

    return walk([el for page in deck.get("slides") or [] for el in page.get("pageElements") or []])


def _slides_shape_style(shape_el: dict) -> tuple[dict, dict]:
    """The shape's first run style and first paragraph style, filtered to writable fields.

    Fields the shape inherits come back unset, which is what we want — reapplying
    only what was explicitly set leaves inheritance intact.
    """
    text = (shape_el.get("shape") or {}).get("text") or {}
    run_style: dict = {}
    para_style: dict = {}
    for item in text.get("textElements") or []:
        if not run_style:
            run_style = (item.get("textRun") or {}).get("style") or {}
        if not para_style:
            para_style = (item.get("paragraphMarker") or {}).get("style") or {}
        if run_style and para_style:
            break
    return (
        {k: v for k, v in run_style.items() if k in _TEXT_STYLE_FIELDS},
        {k: v for k, v in para_style.items() if k in _PARA_STYLE_FIELDS},
    )


def slides_set_shape_text(file_id: str, object_id: str, text: str) -> dict:
    """Replace the text of one shape, identified by the object id slides_read returns.

    Deleting a shape's text throws away its run styling, so the existing font,
    size, colour and paragraph spacing are captured first and reapplied to the
    new text. Without that, every edit silently resets the slide to Slides'
    defaults.
    """
    svc = _build("slides", "v1")
    deck = svc.presentations().get(
        presentationId=file_id, fields="slides(pageElements)"
    ).execute()
    element = _slides_find_element(deck, object_id)
    if element is None:
        raise ValueError(
            f'No shape "{object_id}" in this deck. Read it first — object ids are '
            "per-deck and change when a slide is recreated."
        )
    run_style, para_style = _slides_shape_style(element)

    whole = {"type": "ALL"}
    requests: list[dict] = []
    if _slides_shape_text(element):
        requests.append({"deleteText": {"objectId": object_id, "textRange": whole}})
    requests.append({"insertText": {"objectId": object_id, "insertionIndex": 0, "text": text}})
    if run_style:
        requests.append({"updateTextStyle": {
            "objectId": object_id,
            "textRange": whole,
            "style": run_style,
            "fields": ",".join(sorted(run_style)),
        }})
    if para_style:
        requests.append({"updateParagraphStyle": {
            "objectId": object_id,
            "textRange": whole,
            "style": para_style,
            "fields": ",".join(sorted(para_style)),
        }})

    svc.presentations().batchUpdate(
        presentationId=file_id, body={"requests": requests}
    ).execute()
    return {"object_id": object_id, "chars": len(text), "style_kept": bool(run_style)}


def slides_delete_slide(file_id: str, object_id: str) -> None:
    """Delete one slide by its page object id."""
    svc = _build("slides", "v1")
    svc.presentations().batchUpdate(
        presentationId=file_id,
        body={"requests": [{"deleteObject": {"objectId": object_id}}]},
    ).execute()


def drive_get_metadata(file_id: str) -> dict | None:

    """Fetch lightweight Drive file metadata for update detection.

    Returns ``{id, name, mimeType, modifiedTime, trashed, url}`` or ``None`` if
    the file no longer exists (404) / is otherwise inaccessible.  This is a
    cheap metadata-only call (no content download) suitable for polling.
    """
    from googleapiclient.errors import HttpError

    svc = _build("drive", "v3")
    try:
        meta = svc.files().get(
            fileId=file_id,
            fields="id,name,mimeType,modifiedTime,trashed,webViewLink",
            supportsAllDrives=True,
        ).execute()
    except HttpError as exc:
        if getattr(exc, "resp", None) is not None and exc.resp.status in (404, 403):
            return None
        raise
    return {
        "id": meta.get("id", file_id),
        "name": meta.get("name", ""),
        "mimeType": meta.get("mimeType", ""),
        "modified": meta.get("modifiedTime", ""),
        "trashed": bool(meta.get("trashed", False)),
        "url": meta.get("webViewLink", ""),
    }


# ── Calendar ──────────────────────────────────────────────────────────────

def calendar_events(days_behind: int = 0, days_ahead: int = 7) -> list[dict]:
    svc = _build("calendar", "v3")
    now = datetime.now(timezone.utc)
    resp = svc.events().list(
        calendarId="primary",
        timeMin=(now - timedelta(days=days_behind)).isoformat(),
        timeMax=(now + timedelta(days=days_ahead)).isoformat(),
        maxResults=500, singleEvents=True, orderBy="startTime",
    ).execute()
    out = []
    for e in resp.get("items", []):
        start = e.get("start", {})
        end   = e.get("end", {})
        out.append({
            "id": e.get("id", ""),
            "title": e.get("summary", "(No title)"),
            "start": start.get("dateTime", start.get("date", "")),
            "end":   end.get("dateTime", end.get("date", "")),
            "location": e.get("location", ""),
            "description": (e.get("description") or "")[:500],
            "attendees": [a.get("email", "") for a in e.get("attendees", [])],
            "url": e.get("htmlLink", ""),
        })
    return out


def calendar_create_event(
    title: str, start: str, end: str,
    description: str = "", location: str = "", attendees: list[str] | None = None,
) -> dict:
    svc = _build("calendar", "v3")
    body: dict[str, Any] = {
        "summary": title,
        "start": {"dateTime": start, "timeZone": "UTC"},
        "end":   {"dateTime": end,   "timeZone": "UTC"},
        "description": description,
        "location": location,
        "attendees": [{"email": e} for e in (attendees or [])],
    }
    event = svc.events().insert(calendarId="primary", body=body).execute()
    return {"id": event["id"], "url": event.get("htmlLink", ""), "status": "created"}


# ── Contacts ──────────────────────────────────────────────────────────────

def contacts_search(query: str, max_results: int = 10) -> list[dict]:
    svc = _build("people", "v1")
    resp = svc.people().searchContacts(
        query=query,
        readMask="names,emailAddresses,phoneNumbers,organizations",
        pageSize=max_results,
    ).execute()
    out = []
    for r in resp.get("results", []):
        p = r.get("person", {})
        names  = p.get("names", [])
        emails = p.get("emailAddresses", [])
        phones = p.get("phoneNumbers", [])
        orgs   = p.get("organizations", [])
        out.append({
            "name":    names[0].get("displayName", "")  if names  else "",
            "email":   emails[0].get("value", "")       if emails else "",
            "phone":   phones[0].get("value", "")       if phones else "",
            "company": orgs[0].get("name", "")          if orgs   else "",
        })
    return out


# ── Sheets ────────────────────────────────────────────────────────────────

def sheets_read(spreadsheet_id: str, range_: str = "Sheet1") -> dict:
    """Read values from a Google Sheet range (e.g. 'Sheet1!A1:Z100')."""
    svc = _build("sheets", "v4")
    result = svc.spreadsheets().values().get(
        spreadsheetId=spreadsheet_id,
        range=range_,
    ).execute()
    rows = result.get("values", [])
    return {
        "spreadsheet_id": spreadsheet_id,
        "range": result.get("range", range_),
        "rows": rows,
        "row_count": len(rows),
    }


def sheets_get_metadata(spreadsheet_id: str) -> dict:
    """Get sheet names and basic metadata for a spreadsheet."""
    svc = _build("sheets", "v4")
    meta = svc.spreadsheets().get(
        spreadsheetId=spreadsheet_id,
        fields="spreadsheetId,properties.title,sheets.properties",
    ).execute()
    return {
        "id": meta.get("spreadsheetId", ""),
        "title": meta.get("properties", {}).get("title", ""),
        "sheets": [
            s["properties"]["title"]
            for s in meta.get("sheets", [])
        ],
    }


# ── Sheets write helpers (Little Gerry Budgets) ───────────────────────────
# These only ever touch spreadsheets CREATED by Little Gerry — which is why
# they work under the existing drive.file scope with no re-consent.

def sheets_create_budget_spreadsheet(
    title: str,
    allotment: float | None,
    currency: str,
    categories: list[str],
    folder_id: str | None = None,
) -> dict:
    """Create a standardized budget spreadsheet (Ledger/Categories/Settings
    tabs, headers, and summary formulas), optionally moved into a folder.
    Returns ``{id, url}``."""
    svc = _build("sheets", "v4")

    body = {
        "properties": {"title": f"{title} — Little Gerry Budget"},
        "sheets": [
            {"properties": {"title": "Ledger", "gridProperties": {"frozenRowCount": 1}}},
            {"properties": {"title": "Categories", "gridProperties": {"frozenRowCount": 1}}},
            {"properties": {"title": "Settings"}},
        ],
    }
    created = svc.spreadsheets().create(
        body=body, fields="spreadsheetId,spreadsheetUrl"
    ).execute()
    sid = created["spreadsheetId"]

    from datetime import datetime as _dt

    values_payload = [
        {"range": "Ledger!A1:F1",
         "values": [["Date", "Description", "Category", "Amount", "Source", "Note"]]},
        {"range": "Categories!A1:B1", "values": [["Category", "Cap"]]},
        {"range": "Settings!A1:B7", "values": [
            ["Title", title],
            ["Allotment", allotment if allotment is not None else ""],
            ["Currency", currency],
            ["Created", _dt.now().strftime("%Y-%m-%d")],
            ["Managed By", "Little Gerry"],
            ["Total Spent", "=SUM(Ledger!D2:D)"],
            ["Remaining", "=IF(B2=\"\",\"\",B2-B6)"],
        ]},
    ]
    if categories:
        values_payload.append({
            "range": f"Categories!A2:A{1 + len(categories)}",
            "values": [[c] for c in categories],
        })
    svc.spreadsheets().values().batchUpdate(
        spreadsheetId=sid,
        body={"valueInputOption": "USER_ENTERED", "data": values_payload},
    ).execute()

    url = created.get("spreadsheetUrl", f"https://docs.google.com/spreadsheets/d/{sid}")

    if folder_id:
        drive = _build("drive", "v3")
        try:
            meta = drive.files().get(fileId=sid, fields="parents", supportsAllDrives=True).execute()
            prev = ",".join(meta.get("parents", []))
            drive.files().update(
                fileId=sid,
                addParents=folder_id,
                removeParents=prev,
                fields="id",
                supportsAllDrives=True,
            ).execute()
        except Exception:  # noqa: BLE001 — placement is cosmetic; the sheet works from root
            logger.warning("Budget sheet %s created but could not be moved to folder", sid)

    return {"id": sid, "url": url}


def sheets_append_row(spreadsheet_id: str, range_: str, values: list) -> None:
    """Append one row of values to a table range (e.g. 'Ledger!A:F')."""
    svc = _build("sheets", "v4")
    svc.spreadsheets().values().append(
        spreadsheetId=spreadsheet_id,
        range=range_,
        valueInputOption="USER_ENTERED",
        insertDataOption="INSERT_ROWS",
        body={"values": [values]},
    ).execute()


def sheets_update_range(spreadsheet_id: str, range_: str, values: list[list]) -> None:
    """Overwrite an exact range with values (targeted, never wholesale)."""
    svc = _build("sheets", "v4")
    svc.spreadsheets().values().update(
        spreadsheetId=spreadsheet_id,
        range=range_,
        valueInputOption="USER_ENTERED",
        body={"values": values},
    ).execute()


def sheets_delete_row(spreadsheet_id: str, tab_title: str, row_index_1based: int) -> None:
    """Delete a single row from a named tab (row 1 = first row)."""
    svc = _build("sheets", "v4")
    meta = svc.spreadsheets().get(
        spreadsheetId=spreadsheet_id, fields="sheets.properties"
    ).execute()
    gid = None
    for s in meta.get("sheets", []):
        if s["properties"].get("title") == tab_title:
            gid = s["properties"].get("sheetId")
            break
    if gid is None:
        raise ValueError(f"Tab '{tab_title}' not found in spreadsheet.")
    svc.spreadsheets().batchUpdate(
        spreadsheetId=spreadsheet_id,
        body={"requests": [{
            "deleteDimension": {
                "range": {
                    "sheetId": gid,
                    "dimension": "ROWS",
                    "startIndex": row_index_1based - 1,
                    "endIndex": row_index_1based,
                }
            }
        }]},
    ).execute()


# ── Tasks ─────────────────────────────────────────────────────────────────

def tasks_list(max_results: int = 25, show_completed: bool = False) -> list[dict]:
    """List tasks across all Google Task lists (default: incomplete only)."""
    svc = _build("tasks", "v1")
    lists_resp = svc.tasklists().list(maxResults=10).execute()
    out: list[dict] = []
    for lst in lists_resp.get("items", []):
        tasks_resp = svc.tasks().list(
            tasklist=lst["id"],
            maxResults=max_results,
            showCompleted=show_completed,
            showHidden=False,
        ).execute()
        for t in tasks_resp.get("items", []):
            out.append({
                "id": t["id"],
                "title": t.get("title", ""),
                "list": lst.get("title", ""),
                "due": t.get("due", ""),
                "notes": t.get("notes", ""),
                "status": t.get("status", ""),
            })
    return out
