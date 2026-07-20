"""
Agent tool definitions and implementations.

Tools available to the PMI Executive Assistant:
  - search_knowledge_base  : semantic search over ingested documents
  - create_task            : create a task in the task tracker (auto-approved)
  - request_approval       : create an ApprovalIntent for human review (high-risk actions)
  - get_pending_approvals  : list open approval items for the current user

Each tool has:
  - DEFINITION: the Ollama tool schema (passed in the system message)
  - An async execute_*(ctx, args) function that runs the tool

Tools receive a ToolContext rather than direct DB sessions so the caller
(the agent executor) can inject deps cleanly.
"""

from __future__ import annotations

import logging
import os
import re
import uuid
from dataclasses import dataclass
from datetime import datetime, timedelta, timezone
from pathlib import Path
from typing import Any

from sqlalchemy.ext.asyncio import AsyncSession

from models.db.enums import MessageRole, TaskPriority, TaskStatus
from models.db.task import Task
from repositories.conversation_repo import ApprovalRepository
from repositories.document_repo import DocumentChunkRepository, DocumentRepository
from services.embeddings.service import EmbeddingService

logger = logging.getLogger(__name__)


# ── Tool context ──────────────────────────────────────────────────────────────

@dataclass
class ToolContext:
    db: AsyncSession
    user_id: uuid.UUID
    conversation_id: uuid.UUID
    embedding_service: EmbeddingService
    # When a tool wants the client to show a final confirm/cancel popup before a
    # destructive action, it sets this to a frame dict (e.g. a "confirm_delete"
    # payload). The agent loop emits it to the WebSocket and clears it. The
    # actual deletion happens client-side only after the user confirms.
    pending_confirmation: dict[str, Any] | None = None


# ── Tool schema definitions (sent to Ollama) ──────────────────────────────────

TOOL_DEFINITIONS: list[dict] = [
    {
        "type": "function",
        "function": {
            "name": "search_knowledge_base",
            "description": (
                "Search the local Knowledge Base of *imported* documents. "
                "ONLY call this for PMI documents, VACTOR specifications, regulatory submissions, "
                "protocols, or internal knowledge that has been explicitly uploaded to the KB. "
                "For documents still on Google Drive that have NOT been imported, use search_drive_content instead. "
                "Use this to find specific facts or passages. When the user asks you to summarize, "
                "review, or analyze an ENTIRE document, call read_knowledge_base_document instead so "
                "you see every section — search only returns the few most-similar chunks. "
                "Do NOT call for general questions or things you already know."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "query": {
                        "type": "string",
                        "description": "The search query — a natural language question or topic.",
                    },
                    "top_k": {
                        "type": "integer",
                        "description": "Number of results to return (1–10). Default 5.",
                        "default": 5,
                    },
                },
                "required": ["query"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "read_knowledge_base_document",
            "description": (
                "Read the COMPLETE text of a single imported Knowledge Base document, from start "
                "to finish, including every section. Use this whenever the user asks you to "
                "summarize, review, analyze, proofread, or extract information from a whole KB "
                "document — search_knowledge_base only returns the few most-similar chunks and will "
                "miss sections. Identify the document by document_id (preferred — get it from "
                "manage_knowledge_base or search_knowledge_base) or by a title/query."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "document_id": {
                        "type": "string",
                        "description": "The UUID of the KB document to read (preferred).",
                    },
                    "query": {
                        "type": "string",
                        "description": "A title or search term to identify the document, if the id is unknown.",
                    },
                },
                "required": [],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "request_kb_deletion",
            "description": (
                "Request permanent deletion of a Knowledge Base document. "
                "Use this ONLY when the user explicitly asks to delete or remove a KB document. "
                "This does NOT delete anything itself — it shows the user a confirmation popup where "
                "they give final approval; the document is only removed if they confirm there. "
                "Identify the document by document_id (preferred — get it from manage_knowledge_base "
                "list or search_knowledge_base) or by a title/query. After calling this, stop and wait "
                "for the user's decision; do not call it again for the same document."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "document_id": {
                        "type": "string",
                        "description": "The UUID of the KB document to delete (preferred).",
                    },
                    "query": {
                        "type": "string",
                        "description": "A title or search term to identify the document, if the id is unknown.",
                    },
                },
                "required": [],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "create_task",
            "description": (
                "Create a task in the PMI task tracker. "
                "Use this when the user explicitly asks to create, add, or track a task or action item."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "title": {"type": "string", "description": "Short task title (max 200 chars)."},
                    "description": {"type": "string", "description": "Optional longer description."},
                    "priority": {
                        "type": "string",
                        "enum": ["low", "medium", "high", "critical"],
                        "description": "Task priority.",
                        "default": "medium",
                    },
                    "due_date": {
                        "type": "string",
                        "description": "Optional ISO 8601 due date, e.g. '2026-06-30'.",
                    },
                },
                "required": ["title"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "request_approval",
            "description": (
                "Submit an action for human approval before execution. "
                "MUST be used for any action with real-world consequences: "
                "sending emails, external communications, document modifications, "
                "regulatory submissions, purchases, or anything irreversible."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "intent_type": {
                        "type": "string",
                        "enum": [
                            "send_email",
                            "create_calendar_event",
                            "modify_document",
                            "create_regulatory_submission",
                            "send_message",
                            "delete_record",
                            "external_api_call",
                            "create_task",
                            "update_task",
                        ],
                        "description": "Type of action being requested.",
                    },
                    "title": {"type": "string", "description": "One-line summary of the action."},
                    "description": {
                        "type": "string",
                        "description": "Full description of what will happen if approved.",
                    },
                    "risk_level": {
                        "type": "string",
                        "enum": ["low", "medium", "high", "critical"],
                        "description": "Assessed risk level of the action.",
                        "default": "medium",
                    },
                    "payload": {
                        "type": "object",
                        "description": "Structured data for the action (e.g. email body, recipient).",
                    },
                    "expires_hours": {
                        "type": "integer",
                        "description": "Hours until this approval request expires (default 72).",
                        "default": 72,
                    },
                },
                "required": ["intent_type", "title", "description", "payload"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "create_email_draft",
            "description": (
                "Draft an email and save it to Communications → Email Drafts for the user to "
                "review, edit, and send. Use this whenever the user asks you to DRAFT, WRITE, or "
                "COMPOSE an email — this is the normal case. You write the full email body yourself "
                "and pass it as 'body' (use real line breaks between paragraphs). This does NOT send "
                "anything; it just files a draft the user can open on the Email Drafts page. "
                "RECIPIENT ADDRESS: always try to provide recipient_email. If you only know the "
                "name, the tool automatically looks the address up in the user's contacts — when "
                "exactly one match exists it is used; when none or several match, the tool tells "
                "you and you MUST ask the user which address to use before drafting again. "
                "ATTACHMENTS: pass generated filenames in 'attachments' to attach files from the "
                "Generated Files store. To attach a document that doesn't exist yet, create it "
                "FIRST with create_docx/generate_file, then pass the returned filename here. "
                "Only use request_approval(intent_type='send_email') instead when the user explicitly "
                "asks you to SEND an email right now."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "subject": {"type": "string", "description": "The email subject line."},
                    "body": {
                        "type": "string",
                        "description": (
                            "The full email body you have written, salutation through sign-off, "
                            "with real line breaks (\\n) between paragraphs. End with a simple "
                            "sign-off line (e.g. the user's first name) — do NOT add a signature "
                            "block; the user's configured signature is appended automatically."
                        ),
                    },
                    "recipient_name": {"type": "string", "description": "Recipient's name, if known."},
                    "recipient_email": {"type": "string", "description": "Recipient's email address, if known."},
                    "purpose": {"type": "string", "description": "One line on what the email accomplishes (optional)."},
                    "tone": {
                        "type": "string",
                        "enum": ["professional", "friendly", "formal", "concise", "empathetic", "persuasive"],
                        "description": "Tone of the email.",
                        "default": "professional",
                    },
                    "attachments": {
                        "type": "array",
                        "items": {"type": "string"},
                        "description": (
                            "Filenames from the Generated Files store to attach (as returned by "
                            "create_docx/generate_file, e.g. 'ab12cd34_Report.docx'; the plain "
                            "display name also works). They are attached to the real email when "
                            "the user approves the send."
                        ),
                    },
                },
                "required": ["subject", "body"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "propose_odoo_write",
            "description": (
                "Propose a WRITE to the connected Odoo ERP. This NEVER writes directly — "
                "it queues a pending approval the user must approve. Use for: confirming a "
                "quotation, registering an invoice payment, creating a CRM lead, logging an "
                "internal note on a record, updating fields on a record, or creating a contact."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "action": {
                        "type": "string",
                        "enum": [
                            "confirm_quotation",
                            "register_payment",
                            "create_lead",
                            "log_note",
                            "update_field",
                            "create_contact",
                        ],
                        "description": "Which Odoo write to perform.",
                    },
                    "params": {
                        "type": "object",
                        "description": (
                            "Action parameters. confirm_quotation: {order_id}. "
                            "register_payment: {move_id, amount?}. "
                            "create_lead: {name, contact_name?, email_from?, phone?, expected_revenue?, description?}. "
                            "log_note: {model, record_id, body}. "
                            "update_field: {model, record_id, values:{field:value}}. "
                            "create_contact: {name, email?, phone?, city?}."
                        ),
                    },
                },
                "required": ["action", "params"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "get_pending_approvals",
            "description": "Retrieve a list of pending approval requests awaiting human decision.",
            "parameters": {
                "type": "object",
                "properties": {},
                "required": [],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "get_tasks",
            "description": (
                "Query the PMI task tracker to list tasks for the current user. "
                "Use this to answer questions about what tasks are open, overdue, or due soon."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "status": {
                        "type": "string",
                        "enum": ["backlog", "todo", "in_progress", "in_review", "done", "cancelled", "all"],
                        "description": "Filter by task status. 'all' returns every status.",
                        "default": "all",
                    },
                    "priority": {
                        "type": "string",
                        "enum": ["low", "medium", "high", "critical", "any"],
                        "description": "Filter by priority. 'any' skips priority filtering.",
                        "default": "any",
                    },
                },
                "required": [],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "get_regulatory_status",
            "description": (
                "Retrieve the current regulatory compliance status for VACTOR: "
                "regulatory document counts by status, open/in-progress CAPAs, "
                "and any documents past their review date. "
                "Use for questions about compliance posture, audit readiness, or CAPA status."
            ),
            "parameters": {
                "type": "object",
                "properties": {},
                "required": [],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "search_web",
            "description": (
                "Search the internet using DuckDuckGo for current information, news, "
                "ONLY call this when the user explicitly asks to search the internet or look up "
                "current information online. Do NOT call for questions you can answer directly. "
                "Always cite the URLs you reference."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "query": {
                        "type": "string",
                        "description": "The web search query.",
                    },
                    "max_results": {
                        "type": "integer",
                        "description": "Number of results to return (1–10). Default 5.",
                        "default": 5,
                    },
                },
                "required": ["query"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "fetch_page",
            "description": (
                "Fetch and read the text content of a specific web page URL. "
                "Use this to read the full content of a search result or any public URL."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "url": {
                        "type": "string",
                        "description": "The full URL to fetch.",
                    },
                },
                "required": ["url"],
            },
        },
    },
    # ── Google Workspace tools (read-only — only call when user explicitly asks) ─
    {
        "type": "function",
        "function": {
            "name": "search_gmail",
            "description": (
                "ONLY call this when the user explicitly asks to search, check, or find emails "
                "in their Gmail inbox. Do NOT call for general conversation or greetings. "
                "Uses Gmail search syntax (e.g. 'from:alice subject:VACTOR is:unread')."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "query": {
                        "type": "string",
                        "description": "Gmail search query string.",
                    },
                    "max_results": {
                        "type": "integer",
                        "description": "Max emails to return (1–20). Default 10.",
                        "default": 10,
                    },
                },
                "required": ["query"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "read_gmail_message",
            "description": (
                "Read the full body of a Gmail message by its ID. "
                "Only call after search_gmail has returned results and the user wants to read a specific email."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "message_id": {
                        "type": "string",
                        "description": "The Gmail message ID (from search_gmail results).",
                    },
                },
                "required": ["message_id"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "search_drive",
            "description": (
                "Search for files in Google Drive by keyword. Use this when the user wants to "
                "find a specific file or search across file content. "
                "To list or browse folders, use list_drive_folder instead."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "query": {
                        "type": "string",
                        "description": "Text to search for across file names and content.",
                    },
                    "max_results": {
                        "type": "integer",
                        "description": "Max files to return (1–20). Default 10.",
                        "default": 10,
                    },
                },
                "required": ["query"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "list_drive_folder",
            "description": (
                "List the contents (files and sub-folders) of a Google Drive folder. "
                "Use this when the user asks to list, browse, or see what is in a Drive folder. "
                "Call with no folder_id (or folder_id='root') to list the top-level My Drive. "
                "Use a folder ID from a previous list_drive_folder result to browse deeper. "
                "NOTE: top-level folders may live in a SHARED drive — call list_shared_drives first, "
                "then pass both folder_id and drive_id (= the shared drive's ID) to list its root."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "folder_id": {
                        "type": "string",
                        "description": "Drive folder ID to list. Use 'root' for the top level. Default is 'root'.",
                        "default": "root",
                    },
                    "drive_id": {
                        "type": "string",
                        "description": "Shared drive ID — required when listing the root of a shared drive (use the ID from list_shared_drives for both folder_id and drive_id).",
                    },
                    "max_results": {
                        "type": "integer",
                        "description": "Max items to return (1–50). Default 50.",
                        "default": 50,
                    },
                },
                "required": [],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "list_shared_drives",
            "description": (
                "List all Google shared (team) drives the account can access — these are the "
                "top-level drive trees that sit BESIDE My Drive and are invisible to "
                "list_drive_folder('root'). Use this FIRST when the user asks about top-level "
                "folders (e.g. Communications, Knowledge, Compliance) that don't appear in My Drive, "
                "then use list_drive_folder with the returned drive ID to browse inside."
            ),
            "parameters": {"type": "object", "properties": {}, "required": []},
        },
    },
    {
        "type": "function",
        "function": {
            "name": "search_drive_content",
            "description": (
                "Search Google Drive files by keyword and read their full text content. "
                "Use this to answer questions about documents stored on Google Drive that have NOT "
                "been imported into the Knowledge Base. Searches across file names and content, "
                "then automatically reads the text of matching files. "
                "Use search_knowledge_base for already-imported documents instead."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "query": {
                        "type": "string",
                        "description": "Keywords or phrase to search for across Drive file names and content.",
                    },
                    "max_files": {
                        "type": "integer",
                        "description": "How many matching files to read (1–5). Default 3.",
                        "default": 3,
                    },
                },
                "required": ["query"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "read_drive_file",
            "description": (
                "Read the text content of a Google Drive file — including live Google "
                "Docs the user is currently writing. Call this whenever the user shares "
                "a Drive/Docs/Sheets link or asks for feedback, input, or recommendations "
                "on a document they're working on: read it, then give concrete suggestions. "
                "Works for Docs, Sheets (as CSV), Slides, PDFs, and plain text files. "
                "Long documents are returned in 30,000-character pages — when the result "
                "ends with a CONTINUE note, call again with the suggested offset to read "
                "the next page. NEVER review only the first page of a long document and "
                "present conclusions as complete — page through to the end first."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "file_id": {
                        "type": "string",
                        "description": (
                            "The Drive file ID (from search_drive results) OR a full pasted "
                            "URL like https://docs.google.com/document/d/<id>/edit — the ID "
                            "is extracted automatically."
                        ),
                    },
                    "offset": {
                        "type": "number",
                        "description": (
                            "Character position to start reading from (default 0). Use the "
                            "offset suggested in the CONTINUE note to read the next page of "
                            "a long document."
                        ),
                    },
                },
                "required": ["file_id"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "list_recent_drive_files",
            "description": (
                "List the user's most recently modified Google Drive files (newest "
                "first). Call this when the user says 'help me with this document' or "
                "similar WITHOUT a link — the doc they mean is almost always at the top. "
                "Present the top few and confirm which one, then follow_drive_document."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "max_results": {
                        "type": "integer",
                        "description": "How many files to list (default 8).",
                        "default": 8,
                    },
                },
                "required": [],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "follow_drive_document",
            "description": (
                "Follow a Google Doc/Drive file in THIS conversation for live "
                "collaboration: from then on its CURRENT contents are re-read "
                "automatically on every message, so you always see the user's latest "
                "edits. Use when the user wants ongoing help writing or revising a "
                "document. Accepts a file ID or a pasted URL. Use unfollow_drive_document "
                "to stop."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "file_id": {
                        "type": "string",
                        "description": "Drive file ID or full pasted Docs/Drive URL.",
                    },
                },
                "required": ["file_id"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "unfollow_drive_document",
            "description": "Stop following the live document in this conversation.",
            "parameters": {"type": "object", "properties": {}, "required": []},
        },
    },
    {
        "type": "function",
        "function": {
            "name": "add_to_knowledge_base",
            "description": (
                "Add a file to the Knowledge Base so it becomes searchable and citable. "
                "Use when the user asks to save/import/add a file to the KB. Two sources: "
                "a Google Drive file (pass drive_file_id — an ID or pasted URL; use "
                "search_drive or list_recent_drive_files to find it) or a generated file "
                "(pass generated_filename from Generated Files). Duplicates are detected "
                "and reported. Never imports as a regulated document — the user does that "
                "manually on the Knowledge Base page."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "drive_file_id": {
                        "type": "string",
                        "description": "Google Drive file ID or full pasted URL (for Drive files).",
                    },
                    "generated_filename": {
                        "type": "string",
                        "description": "Exact filename of a generated file (for files Gerry created).",
                    },
                    "title": {
                        "type": "string",
                        "description": "Optional Knowledge Base title (defaults to the file name).",
                    },
                    "category": {
                        "type": "string",
                        "description": "Optional category name (created if missing).",
                    },
                },
                "required": [],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "check_drive_backup_status",
            "description": (
                "Verify the company Google Drive backup is current. Reads the GCS "
                "backup bucket (read-only) for the last backup write, object count and "
                "size, compares against the live Shared Drive, and reports a "
                "CURRENT/STALE verdict plus which files changed since the last backup. "
                "Use when the user asks whether the Drive is backed up or how fresh "
                "the backup is."
            ),
            "parameters": {"type": "object", "properties": {}, "required": []},
        },
    },
    {
        "type": "function",
        "function": {
            "name": "get_calendar_events",
            "description": (
                "ONLY call this when the user explicitly asks to check their calendar, "
                "see upcoming meetings, or asks what is scheduled. "
                "Do NOT call proactively or for general greetings."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "days_behind": {
                        "type": "integer",
                        "description": "Days before today to include (default 0).",
                        "default": 0,
                    },
                    "days_ahead": {
                        "type": "integer",
                        "description": "Days ahead to include (default 7).",
                        "default": 7,
                    },
                },
                "required": [],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "search_contacts",
            "description": (
                "Look up a person's contact details (email, phone, company). Use when "
                "the user asks about a contact or 'who is our contact at <company>'. "
                "Searches both PMI's own contacts (derived from email + manual entries) "
                "and Google Contacts."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "query": {
                        "type": "string",
                        "description": "Name, email address, or company to search for.",
                    },
                },
                "required": ["query"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "add_contacts",
            "description": (
                "Add one or more contacts to PMI's own Contacts page (Communications → "
                "Contacts — Gerry's local contact store, NOT Odoo and NOT Google "
                "Contacts). Use when the user asks to save, import, or add people to "
                "their contacts. Existing entries with the same email are updated. "
                "For Odoo CRM contacts use propose_odoo_write instead."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "contacts": {
                        "type": "array",
                        "description": "The contacts to add.",
                        "items": {
                            "type": "object",
                            "properties": {
                                "email": {"type": "string", "description": "Email address (required)."},
                                "name": {"type": "string", "description": "Full name."},
                                "company": {"type": "string", "description": "Company or organization."},
                                "notes": {"type": "string", "description": "Free-text notes (role, phone, context)."},
                            },
                            "required": ["email"],
                        },
                    },
                },
                "required": ["contacts"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "read_google_sheet",
            "description": (
                "ONLY call this when the user explicitly provides a spreadsheet ID or asks "
                "to read data from a specific Google Sheet. "
                "Requires a spreadsheet ID (found in the URL between /d/ and /edit)."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "spreadsheet_id": {
                        "type": "string",
                        "description": "The Google Sheets spreadsheet ID (from the URL between /d/ and /edit).",
                    },
                    "range": {
                        "type": "string",
                        "description": "Cell range to read, e.g. 'Sheet1' or 'Sheet1!A1:Z100'. Default 'Sheet1'.",
                        "default": "Sheet1",
                    },
                },
                "required": ["spreadsheet_id"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "list_google_tasks",
            "description": (
                "ONLY call this when the user explicitly asks to see their Google Tasks "
                "or tasks from Google Workspace. Do NOT call for PMI internal tasks — "
                "use get_tasks for those."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "max_results": {
                        "type": "integer",
                        "description": "Max tasks to return per list (default 25).",
                        "default": 25,
                    },
                },
                "required": [],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "get_file_template",
            "description": (
                "Fetch the company's required structure/format for a document type "
                "from the shared templates folder on Drive. ALWAYS call this before "
                "creating a document with generate_file or create_docx (types like "
                "memo, SOP, letter, report, meeting-minutes). If a template exists, "
                "follow it exactly. Even when no template matches, the company STYLE "
                "GUIDE is returned — apply it so every document looks uniform."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "file_type": {
                        "type": "string",
                        "description": "Document type, e.g. 'memo', 'SOP', 'letter', 'report'.",
                    },
                },
                "required": ["file_type"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "create_workroom",
            "description": (
                "Create a new Workroom — a persistent co-work space with a goal, "
                "pinned artifacts, a dedicated conversation, and a progress "
                "journal. Use when the user asks to set up a room / workspace "
                "for an ongoing effort (a submission, an audit, a fundraise). "
                "After creating, you can pin items with add_to_workroom."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "title": {
                        "type": "string",
                        "description": "Room title, e.g. '510(k) Submission'.",
                    },
                    "goal": {
                        "type": "string",
                        "description": "What the room is working toward (one or two sentences).",
                    },
                },
                "required": ["title"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "add_to_workroom",
            "description": (
                "Pin an artifact to a Workroom (persistent co-work space) so it is "
                "carried into every future conversation turn in that room. Use when "
                "the user says to pin/add/keep something in the room, or when an "
                "artifact is clearly central to the room's goal. In a room "
                "conversation the current room is used automatically; otherwise (or "
                "to target another room) pass workroom_title."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "kind": {
                        "type": "string",
                        "enum": [
                            "drive_doc", "kb_doc", "generated_file", "note",
                            "email_thread", "task", "odoo_record", "regulatory_doc",
                        ],
                        "description": "Artifact kind. Use 'note' for plain facts/decisions worth keeping.",
                    },
                    "label": {
                        "type": "string",
                        "description": "Short human-readable label, e.g. 'QMS Manual draft'. For notes, the note text itself.",
                    },
                    "ref_id": {
                        "type": "string",
                        "description": "Reference: Drive file ID, KB document id, generated filename, Gmail thread id, task id, etc. Omit for notes.",
                    },
                    "workroom_title": {
                        "type": "string",
                        "description": "Target room title (fuzzy matched). Omit inside a room conversation.",
                    },
                },
                "required": ["kind", "label"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "list_workroom_items",
            "description": (
                "List a Workroom's goal, pinned artifacts, and recent journal "
                "entries. In a room conversation the current room is used "
                "automatically; otherwise pass workroom_title (or omit to see "
                "which rooms exist)."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "workroom_title": {
                        "type": "string",
                        "description": "Room title (fuzzy matched). Omit inside a room conversation.",
                    },
                },
                "required": [],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "log_workroom_progress",
            "description": (
                "Append a progress entry to a Workroom's journal — a shared "
                "timeline of what you and the user have accomplished. Log when a "
                "milestone is reached, a decision is made, or the user asks you to "
                "note progress. Keep entries to one crisp sentence. In a room "
                "conversation the current room is used automatically."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "entry": {
                        "type": "string",
                        "description": "One-sentence progress note, e.g. 'Finished section 4 of the risk analysis'.",
                    },
                    "workroom_title": {
                        "type": "string",
                        "description": "Target room title (fuzzy matched). Omit inside a room conversation.",
                    },
                },
                "required": ["entry"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "generate_file",
            "description": (
                "Generate a downloadable file and save it to the server. "
                "Use this when the user asks you to create a report, export data, "
                "write a document, or produce any output they can download. "
                "Before writing a business document, call get_file_template with the "
                "document type and follow the returned structure. "
                "Returns a download URL the user can click."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "filename": {
                        "type": "string",
                        "description": "Filename with extension, e.g. 'report.md' or 'tasks.csv'. Allowed: .txt .md .csv .json",
                    },
                    "content": {
                        "type": "string",
                        "description": "Full text content to write into the file.",
                    },
                },
                "required": ["filename", "content"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "create_docx",
            "description": (
                "Create a Microsoft Word (.docx) document and save it to the server's "
                "Generated Files area so the user can download it. Use this whenever the "
                "user asks for a Word document, a formatted report, a memo, a weekly update, "
                "meeting notes, or any deliverable that should be a .docx. "
                "Before writing, call get_file_template with the document type and follow "
                "the returned structure. When the template or style guide specifies page "
                "headers/footers, fonts, or table colors, pass them via the optional "
                "layout fields — they render as REAL Word headers/footers/styling. The "
                "'content' field accepts lightweight Markdown: '# '/'## '/'### ' headings; "
                "'- ' or '* ' bullets; '1. ' numbered items; '**bold**'; blank lines "
                "separate paragraphs; and pipe tables ('| A | B |' rows). Put a "
                "'| --- | --- |' separator after the first row to style it as a colored "
                "header row with alternating shading; OMIT the separator for label/value "
                "grids (metadata blocks) — their label columns (1st, 3rd) render "
                "accent-filled with bold white text and thin borders. "
                "Returns a download URL."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "filename": {
                        "type": "string",
                        "description": "Filename, e.g. 'Weekly Update.docx'. The .docx extension is added if missing.",
                    },
                    "title": {
                        "type": "string",
                        "description": "Optional document title rendered as a heading at the top.",
                    },
                    "content": {
                        "type": "string",
                        "description": "Document body as lightweight Markdown (see tool description for supported syntax).",
                    },
                    "font": {
                        "type": "string",
                        "description": "Optional base font name for body text, e.g. 'Calibri' or 'Arial'.",
                    },
                    "font_size": {
                        "type": "number",
                        "description": "Optional base font size in points, e.g. 11.",
                    },
                    "header_left": {
                        "type": "string",
                        "description": "Optional page-header left text, e.g. the document title.",
                    },
                    "header_right": {
                        "type": "string",
                        "description": "Optional page-header right text, e.g. 'SOP-011 | v0.1'.",
                    },
                    "footer_left": {
                        "type": "string",
                        "description": "Optional page-footer left text, e.g. 'CONFIDENTIAL – INTERNAL USE ONLY'. 'Page X of Y' is added on the right automatically.",
                    },
                    "accent_color": {
                        "type": "string",
                        "description": "Optional hex color (no #) for banners, table header rows and label cells, e.g. '0A2F41' (PMI navy) or '064E44' (PMI teal).",
                    },
                    "banner_label": {
                        "type": "string",
                        "description": "Optional title-block banner: small top line, e.g. 'QMS DOCUMENTATION'. Renders a full-width colored block at the top of page 1.",
                    },
                    "banner_title": {
                        "type": "string",
                        "description": "Banner main title (large bold white), e.g. 'SOP-011 — Supplier Control'. When using the banner, do NOT repeat the title in content.",
                    },
                    "banner_subtitle": {
                        "type": "string",
                        "description": "Banner subtitle line, e.g. 'Standard Operating Procedure'.",
                    },
                },
                "required": ["filename", "content"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "upload_to_drive",
            "description": (
                "Upload a previously generated file (from the Generated Files area) to the "
                "user's Google Drive. Provide the server filename returned by create_docx or "
                "generate_file (e.g. the value after '/api/files/'). Optionally place it inside "
                "a Drive folder by id. Returns the Google Drive link. Requires the user to have "
                "connected Google with write access."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "filename": {
                        "type": "string",
                        "description": "The server filename in the Generated Files area, e.g. 'a1b2c3d4_Weekly Update.docx'.",
                    },
                    "drive_name": {
                        "type": "string",
                        "description": "Optional name to give the file in Google Drive. Defaults to the original filename without the internal id prefix.",
                    },
                    "folder_id": {
                        "type": "string",
                        "description": "Optional Google Drive folder id to upload into. Omit to upload to the Drive root.",
                    },
                },
                "required": ["filename"],
            },
        },
    },
]


# ── Tool implementations ───────────────────────────────────────────────────────

async def execute_search_knowledge_base(ctx: ToolContext, args: dict[str, Any]) -> str:
    query = str(args.get("query", "")).strip()
    top_k = min(int(args.get("top_k", 5)), 10)

    if not query:
        return "Error: query must not be empty."

    try:
        embedding = await ctx.embedding_service.embed(query)
    except Exception as exc:
        return f"Knowledge base search unavailable: {exc}"

    chunk_repo = DocumentChunkRepository(ctx.db)
    results = await chunk_repo.vector_search(embedding, top_k=top_k)

    if not results:
        return "No relevant documents found for that query."

    # Fetch document titles for the returned chunks
    from sqlalchemy import select
    from models.db.document import Document

    doc_ids = list({chunk.document_id for chunk, _ in results})
    doc_rows = await ctx.db.execute(
        select(Document.id, Document.title).where(Document.id.in_(doc_ids))
    )
    doc_title_map = {row.id: row.title for row in doc_rows}

    lines = [f"Found {len(results)} relevant document chunk(s):\n"]
    for i, (chunk, similarity) in enumerate(results, 1):
        title = doc_title_map.get(chunk.document_id, "Unknown document")
        score_pct = round(similarity * 100, 1)
        lines.append(
            f"[{i}] \"{title}\" "
            f"(chunk {chunk.chunk_index}, page {chunk.page_number}, score {score_pct}%)\n"
            f"{chunk.content}"
        )
    lines.append(
        "\n(To read one of these documents in full — every section — call "
        "read_knowledge_base_document with its document_id.)"
    )
    return "\n\n".join(lines)


async def execute_read_knowledge_base_document(ctx: ToolContext, args: dict[str, Any]) -> str:
    """Return the FULL text of a single Knowledge Base document, start to finish.

    Unlike search_knowledge_base (which returns only the few most-similar chunks),
    this reads the entire document so the model can summarize, review, or analyze
    it without missing any sections. Resolve by document_id (preferred) or a
    title/query.
    """
    from pathlib import Path
    from sqlalchemy import select, func
    from models.db.document import Document

    MAX_CHARS = 120_000

    doc_repo = DocumentRepository(ctx.db)
    chunk_repo = DocumentChunkRepository(ctx.db)

    doc = None
    raw_id = args.get("document_id")
    if raw_id:
        try:
            doc = await doc_repo.get_active(uuid.UUID(str(raw_id)))
        except (ValueError, AttributeError):
            doc = None

    query = str(args.get("query", "")).strip()
    if doc is None and query:
        # Try a case-insensitive partial title match first.
        rows = await ctx.db.execute(
            select(Document)
            .where(Document.deleted_at.is_(None))
            .where(func.lower(Document.title).like(f"%{query.lower()}%"))
            .limit(1)
        )
        doc = rows.scalars().first()
        # Fall back to semantic search to locate the best-matching document.
        if doc is None:
            try:
                embedding = await ctx.embedding_service.embed(query)
                hits = await chunk_repo.vector_search(embedding, top_k=1)
                if hits:
                    doc = await doc_repo.get_active(hits[0][0].document_id)
            except Exception:
                doc = None

    if doc is None:
        return (
            "No matching Knowledge Base document found. Provide a valid document_id "
            "(from manage_knowledge_base) or a clearer title."
        )

    # Prefer the exact stored file (preserves headings/structure); fall back to chunks.
    full_text = ""
    try:
        from services.documents.ingestion import _decrypt_file, _extract_text

        extension = Path(doc.file_name or "").suffix.lower() or ".bin"
        raw = _decrypt_file(doc.id, extension)
        full_text = _extract_text(raw, doc.mime_type)
    except Exception:
        full_text = ""

    if not full_text.strip():
        # Reconstruct from stored chunks, removing the token overlap between them.
        from services.documents.ingestion import CHUNK_OVERLAP, WORDS_PER_TOKEN

        overlap_words = int(CHUNK_OVERLAP / WORDS_PER_TOKEN)
        chunks = await chunk_repo.get_by_document(doc.id)
        words: list[str] = []
        for i, ch in enumerate(chunks):
            cw = ch.content.split()
            words.extend(cw if i == 0 else cw[overlap_words:])
        full_text = " ".join(words)

    if not full_text.strip():
        return f'Document "{doc.title}" has no readable text content.'

    truncated = len(full_text) > MAX_CHARS
    body = full_text[:MAX_CHARS]
    header = f'Full text of Knowledge Base document "{doc.title}" (id {doc.id}):\n\n'
    footer = (
        "\n\n[Note: this document is very long and was truncated here. Use "
        "search_knowledge_base for specific details beyond this point.]"
        if truncated
        else ""
    )
    return header + body + footer


async def execute_request_kb_deletion(ctx: ToolContext, args: dict[str, Any]) -> str:
    """Stage a confirm/cancel popup for deleting a KB document.

    This NEVER deletes anything server-side. It resolves the target document and
    sets ``ctx.pending_confirmation`` to a ``confirm_delete`` frame, which the
    agent loop emits to the client. The deletion is performed by the frontend
    only after the user clicks Confirm in the popup.
    """
    from sqlalchemy import select
    from models.db.document import Document

    repo = DocumentRepository(ctx.db)

    doc = None
    raw_id = args.get("document_id")
    if raw_id:
        try:
            doc_id = uuid.UUID(str(raw_id))
        except (ValueError, TypeError):
            return "Error: document_id is not a valid id. Use manage_knowledge_base list to find the correct id."
        doc = await repo.get_active(doc_id)
        if doc is None:
            return "No active knowledge base document found with that id. It may have already been deleted."
    else:
        query = str(args.get("query") or args.get("title") or "").strip()
        if not query:
            return "Error: provide a document_id (preferred) or a title/query to identify the document to delete."
        rows = (
            await ctx.db.execute(
                select(Document)
                .where(Document.deleted_at.is_(None), Document.title.ilike(f"%{query}%"))
                .order_by(Document.created_at.desc())
                .limit(10)
            )
        ).scalars().all()
        if not rows:
            return f'No knowledge base document matches "{query}". Nothing was deleted.'
        if len(rows) > 1:
            options = "\n".join(f"- {d.id} | {d.title}" for d in rows)
            return (
                "Several documents match. Ask the user which one they mean, then call "
                "request_kb_deletion again with the exact document_id:\n" + options
            )
        doc = rows[0]

    ctx.pending_confirmation = {
        "type": "confirm_delete",
        "target": "kb_document",
        "document_id": str(doc.id),
        "title": doc.title,
    }
    return (
        f'A confirmation popup is now shown to the user asking them to permanently delete '
        f'"{doc.title}" from the knowledge base. The document is only removed if they confirm '
        f"there. Stop and wait for their decision — do not call this tool again for this document."
    )


async def execute_create_task(ctx: ToolContext, args: dict[str, Any]) -> str:
    title = str(args.get("title", "")).strip()[:500]
    if not title:
        return "Error: task title must not be empty."

    description = args.get("description")
    priority_raw = str(args.get("priority", "medium")).lower()
    priority = priority_raw if priority_raw in ("low", "medium", "high", "critical") else "medium"

    due_date: datetime | None = None
    if raw_due := args.get("due_date"):
        try:
            due_date = datetime.fromisoformat(str(raw_due)).replace(tzinfo=timezone.utc)
        except ValueError:
            pass

    task = Task(
        title=title,
        description=description,
        status=TaskStatus.TODO,
        priority=TaskPriority(priority),
        due_date=due_date,
        source_conversation_id=ctx.conversation_id,
        created_by=ctx.user_id,
        assignee_id=ctx.user_id,
    )
    ctx.db.add(task)
    await ctx.db.flush()
    await ctx.db.refresh(task)

    due_str = f", due {due_date.date()}" if due_date else ""
    return (
        f"Task created: \"{task.title}\" "
        f"[priority={priority}{due_str}, id={task.id}]"
    )


async def _resolve_recipient_email(ctx: ToolContext, name: str) -> tuple[str | None, list[str]]:
    """Look up a recipient's email by name in Gerry's contacts + Google Contacts.

    Returns ``(email, candidates)``: a single confident match, or the list of
    ambiguous candidates (as "Name <email>" strings) when 0 or >1 match.
    """
    import asyncio
    from services.email_contacts import get_contacts, search_contacts_store
    from services.google_service import contacts_search, get_credentials

    found: dict[str, str] = {}  # email -> display name
    try:
        store = await get_contacts(ctx.db)
        for c in search_contacts_store(store, name, 5):
            email = (c.get("email") or "").lower()
            if email:
                found.setdefault(email, c.get("name") or "")
    except Exception:
        pass
    if get_credentials():
        try:
            google = await asyncio.get_event_loop().run_in_executor(
                None, lambda: contacts_search(name, 5)
            )
            for c in google or []:
                email = (c.get("email") or "").lower()
                if email:
                    found.setdefault(email, c.get("name") or "")
        except Exception:
            pass

    candidates = [f"{n} <{e}>" if n else e for e, n in found.items()]
    if len(found) == 1:
        return next(iter(found)), candidates
    return None, candidates


async def execute_create_email_draft(ctx: ToolContext, args: dict[str, Any]) -> str:
    from models.db.email_draft import EmailDraft

    _TONES = {"professional", "friendly", "formal", "concise", "empathetic", "persuasive"}

    subject = str(args.get("subject", "")).strip()[:500]
    body = str(args.get("body", "")).strip()
    if not subject:
        return "Error: email subject must not be empty."
    if not body:
        return "Error: email body must not be empty — write the full email and pass it as 'body'."

    recipient_name = str(args.get("recipient_name", "")).strip() or None
    recipient_email = str(args.get("recipient_email", "")).strip() or None
    resolved_note = ""

    # The user's configured signature (Settings → Inbox → Signature: gmail /
    # custom / none) is appended to every Gerry draft.
    try:
        from services.email_signature import apply_signature, resolve_signature
        body = apply_signature(body, await resolve_signature(ctx.db))
    except Exception:
        pass  # signature is best-effort — never block the draft

    # No address given — resolve it from contacts when it's obvious; otherwise
    # the model must ask the user rather than filing an unsendable draft.
    if not recipient_email and recipient_name:
        match, candidates = await _resolve_recipient_email(ctx, recipient_name)
        if match:
            recipient_email = match
            resolved_note = f" (address {match} found in contacts)"
        elif candidates:
            return (
                f"Cannot draft yet: several contacts match \"{recipient_name}\": "
                + "; ".join(candidates[:5])
                + ". Ask the user which address to use, then call create_email_draft "
                "again with that recipient_email."
            )
        else:
            return (
                f"Cannot draft yet: no email address found for \"{recipient_name}\" in "
                "the user's contacts. Ask the user for the recipient's email address, "
                "then call create_email_draft again with recipient_email filled in."
            )

    purpose = (str(args.get("purpose", "")).strip() or subject)[:2000]
    tone_raw = str(args.get("tone", "professional")).strip().lower()
    tone = tone_raw if tone_raw in _TONES else "professional"

    # Attachments — resolve names against the Generated Files store. Accept the
    # exact safe name or a plain display name (matched by suffix, unique only).
    att_arg = args.get("attachments") or []
    if isinstance(att_arg, str):
        att_arg = [att_arg]
    attachments: list[dict] = []
    for raw_name in att_arg[:10]:
        want = str(raw_name).strip().lstrip("/")
        if want.startswith("api/files/"):
            want = want[len("api/files/"):]
        if not want:
            continue
        path = (_GENERATED_FILES_DIR / want).resolve()
        if str(path).startswith(str(_GENERATED_FILES_DIR.resolve())) and path.is_file():
            safe_name = path.name
        else:
            candidates = [
                p.name
                for p in _GENERATED_FILES_DIR.glob("*")
                if p.is_file() and p.name.lower().endswith("_" + want.lower())
            ] if _GENERATED_FILES_DIR.is_dir() else []
            if len(candidates) == 1:
                safe_name = candidates[0]
            elif len(candidates) > 1:
                return (
                    f"Cannot attach '{want}': several generated files match "
                    f"({', '.join(candidates[:5])}). Pass the exact filename."
                )
            else:
                return (
                    f"Cannot attach '{want}': no such file in Generated Files. "
                    "Create it first with create_docx/generate_file, then pass the "
                    "filename it returns."
                )
        display = re.sub(r"^[0-9a-f]{8}_", "", safe_name)
        attachments.append({"filename": safe_name, "display_name": display})

    draft = EmailDraft(
        id=uuid.uuid4(),
        subject=subject,
        recipient_name=recipient_name,
        recipient_email=recipient_email,
        purpose=purpose,
        tone=tone,
        key_points=None,
        draft_body=body,
        status="draft",
        tags=["assistant"],
        attachments=attachments,
        created_by=ctx.user_id,
    )
    ctx.db.add(draft)
    await ctx.db.flush()
    await ctx.db.refresh(draft)

    from services.workroom_context import log_room_event
    await log_room_event(
        ctx.db, ctx.conversation_id,
        f"Gerry drafted email: \"{subject}\""
        + (f" to {recipient_name or recipient_email}" if (recipient_name or recipient_email) else ""),
    )

    to_str = ""
    if recipient_name or recipient_email:
        to_str = f" to {recipient_name or recipient_email}"
        if recipient_email and recipient_name:
            to_str += f" <{recipient_email}>"
    att_str = (
        " with attachment(s): " + ", ".join(a["display_name"] for a in attachments)
        if attachments
        else ""
    )
    return (
        f"Email draft saved to Communications → Email Drafts: \"{subject}\"{to_str}{resolved_note}{att_str} "
        f"[id={draft.id}, status=draft]. The user can review, edit, and send it from the "
        "Email Drafts page."
    )


async def execute_request_approval(ctx: ToolContext, args: dict[str, Any]) -> str:
    intent_type = str(args.get("intent_type", "external_api_call"))
    title = str(args.get("title", "")).strip()[:500]
    description = str(args.get("description", "")).strip()
    risk_level = str(args.get("risk_level", "medium")).lower()
    payload = args.get("payload", {})
    expires_hours = min(int(args.get("expires_hours", 72)), 720)

    if not title:
        return "Error: approval title must not be empty."

    expires_at = datetime.now(timezone.utc) + timedelta(hours=expires_hours)

    repo = ApprovalRepository(ctx.db)
    intent = await repo.create(
        user_id=ctx.user_id,
        intent_type=intent_type,
        intent_title=title,
        intent_description=description,
        intent_payload={
            **(payload if isinstance(payload, dict) else {"data": payload}),
            "conversation_id": str(ctx.conversation_id),
        },
        risk_level=risk_level,
        expires_at=expires_at,
    )

    return (
        f"Approval request submitted: \"{title}\" "
        f"[risk={risk_level}, id={intent.id}, expires in {expires_hours}h]. "
        "Waiting for human review before this action can proceed."
    )


async def execute_propose_odoo_write(ctx: ToolContext, args: dict[str, Any]) -> str:
    from sqlalchemy import select
    from models.db.odoo import OdooConnection
    from services import odoo_service as odoo

    action = str(args.get("action", "")).strip()
    params = args.get("params") or {}
    if not isinstance(params, dict):
        return "Error: 'params' must be an object."

    conn = (await ctx.db.execute(
        select(OdooConnection).where(OdooConnection.user_id == ctx.user_id)
    )).scalar_one_or_none()
    if conn is None:
        return "Error: Odoo is not connected. Ask the user to connect it on the Odoo page first."

    try:
        odoo.validate_write(action, params)
    except odoo.OdooError as exc:
        return f"Error: {exc}"

    title, description = odoo.describe_write(action, params)
    intent = await ApprovalRepository(ctx.db).create(
        user_id=ctx.user_id,
        intent_type="odoo_write",
        intent_title=title[:500],
        intent_description=description,
        intent_payload={
            "action": action,
            "params": params,
            "conversation_id": str(ctx.conversation_id),
        },
        risk_level=odoo.default_risk(action),
        expires_at=datetime.now(timezone.utc) + timedelta(hours=72),
    )
    return (
        f"Odoo action queued for approval: \"{title}\" "
        f"[risk={intent.risk_level}, id={intent.id}]. "
        "It will run only after the user approves it on the Approvals page."
    )


async def execute_get_pending_approvals(ctx: ToolContext, _args: dict[str, Any]) -> str:
    repo = ApprovalRepository(ctx.db)
    items = await repo.list_pending(ctx.user_id, limit=10)
    if not items:
        return "No pending approvals."
    lines = [f"Pending approvals ({len(items)}):"]
    for item in items:
        lines.append(
            f"- [{item.risk_level.upper()}] \"{item.intent_title}\" "
            f"(type={item.intent_type}, id={item.id})"
        )
    return "\n".join(lines)


async def execute_get_tasks(ctx: ToolContext, args: dict[str, Any]) -> str:
    from sqlalchemy import select
    from models.db.task import Task
    from models.db.enums import TaskStatus

    status_filter = str(args.get("status", "all"))
    priority_filter = str(args.get("priority", "any"))

    stmt = select(Task).where(
        (Task.created_by == ctx.user_id) | (Task.assignee_id == ctx.user_id)
    )
    if status_filter != "all":
        stmt = stmt.where(Task.status == status_filter)
    if priority_filter != "any":
        stmt = stmt.where(Task.priority == priority_filter)
    stmt = stmt.order_by(Task.due_date.asc().nullslast(), Task.created_at.desc()).limit(30)

    result = await ctx.db.execute(stmt)
    tasks = list(result.scalars())

    if not tasks:
        return "No tasks found matching that filter."

    now = datetime.now(timezone.utc)
    lines = [f"Tasks ({len(tasks)} found):"]
    for t in tasks:
        due_str = ""
        if t.due_date:
            is_overdue = t.due_date < now and t.status not in (TaskStatus.DONE, TaskStatus.CANCELLED)
            due_str = f", due {t.due_date.date()}" + (" [OVERDUE]" if is_overdue else "")
        lines.append(
            f"- [{t.status.upper()}][{t.priority}] {t.title}{due_str}"
        )
    return "\n".join(lines)


async def execute_get_regulatory_status(ctx: ToolContext, _args: dict[str, Any]) -> str:
    from sqlalchemy import func, select
    from models.db.regulatory import RegulatoryDocument
    from models.db.enums import RegDocStatus
    from repositories.regulatory_repo import CAPARepository, RegulatoryDocRepository

    reg_repo = RegulatoryDocRepository(ctx.db)
    capa_repo = CAPARepository(ctx.db)

    all_docs = await reg_repo.list()
    all_capas = await capa_repo.list()

    # Doc counts by status
    status_counts: dict[str, int] = {}
    for doc in all_docs:
        status_counts[doc.status] = status_counts.get(doc.status, 0) + 1

    # Overdue reviews
    from datetime import date
    today = date.today()
    overdue_reviews = [
        d for d in all_docs
        if d.next_review_date and d.next_review_date < today
        and d.status not in ("superseded",)
    ]

    # Open CAPAs
    open_capas = [c for c in all_capas if c.status in ("open", "in_progress")]

    lines = ["VACTOR Regulatory Compliance Status\n"]
    lines.append(f"Total documents: {len(all_docs)}")
    for s, n in sorted(status_counts.items()):
        lines.append(f"  {s}: {n}")

    if overdue_reviews:
        lines.append(f"\nDocuments past review date ({len(overdue_reviews)}):")
        for d in overdue_reviews[:5]:
            lines.append(f"  - {d.title} (rev {d.revision}, review was due {d.next_review_date})")

    lines.append(f"\nCAPAs: {len(all_capas)} total, {len(open_capas)} open/in-progress")
    for c in open_capas[:5]:
        due = f", due {c.due_date}" if getattr(c, "due_date", None) else ""
        lines.append(f"  - [{c.status.upper()}] {c.capa_number}: {c.title}{due}")

    return "\n".join(lines)


async def execute_search_web(ctx: ToolContext, args: dict[str, Any]) -> str:
    from services.research.searcher import web_search

    query = str(args.get("query", "")).strip()
    if not query:
        return "Error: query must not be empty."
    max_results = min(int(args.get("max_results", 5)), 10)

    results = await web_search(query, max_results=max_results)
    if not results:
        return f"No web search results found for: {query}"

    lines = [f"Web search results for \"{query}\" ({len(results)} found):\n"]
    for i, r in enumerate(results, 1):
        lines.append(
            f"[{i}] {r['title']}\n"
            f"    URL: {r['url']}\n"
            f"    {r['snippet']}"
        )
    return "\n\n".join(lines)


async def execute_fetch_page(ctx: ToolContext, args: dict[str, Any]) -> str:
    from services.research.searcher import fetch_page_text

    url = str(args.get("url", "")).strip()
    if not url or not url.startswith("http"):
        return "Error: a valid http/https URL is required."

    text = await fetch_page_text(url, max_chars=4000)
    if not text:
        return f"Could not fetch content from {url} (network error or empty page)."
    return f"Content from {url}:\n\n{text}"


# ── Google Workspace tool executors ───────────────────────────────────────────

def _google_not_connected() -> str:
    return (
        "Google account is not connected. "
        "Ask the user to connect Google via Settings → Google Integration."
    )


async def execute_search_gmail(ctx: ToolContext, args: dict[str, Any]) -> str:
    import asyncio
    from services.google_service import gmail_search, get_credentials
    if not get_credentials():
        return _google_not_connected()
    query = str(args.get("query", "")).strip()
    if not query:
        return "Error: query must not be empty."
    max_results = min(int(args.get("max_results", 10)), 20)
    try:
        msgs = await asyncio.get_event_loop().run_in_executor(
            None, lambda: gmail_search(query, max_results)
        )
    except Exception as exc:
        return f"Gmail search failed: {exc}"
    if not msgs:
        return f"No emails found matching: {query}"
    lines = [f"Gmail results for '{query}' ({len(msgs)} found):\n"]
    for m in msgs:
        lines.append(
            f"ID: {m['id']}\n"
            f"From: {m['from']}\nTo: {m['to']}\n"
            f"Subject: {m['subject']}\nDate: {m['date']}\n"
            f"Snippet: {m['snippet']}"
        )
    return "\n\n".join(lines)


async def execute_read_gmail_message(ctx: ToolContext, args: dict[str, Any]) -> str:
    import asyncio
    from services.google_service import gmail_get_message, get_credentials
    if not get_credentials():
        return _google_not_connected()
    message_id = str(args.get("message_id", "")).strip()
    if not message_id:
        return "Error: message_id is required."
    try:
        msg = await asyncio.get_event_loop().run_in_executor(
            None, lambda: gmail_get_message(message_id)
        )
    except Exception as exc:
        return f"Failed to read email: {exc}"
    return (
        f"From: {msg['from']}\nTo: {msg['to']}\n"
        f"Subject: {msg['subject']}\nDate: {msg['date']}\n\n"
        f"{msg['body'][:6000]}"
    )


async def execute_search_drive(ctx: ToolContext, args: dict[str, Any]) -> str:
    import asyncio
    from services.google_service import drive_search, get_credentials
    if not get_credentials():
        return _google_not_connected()
    query = str(args.get("query", "")).strip()
    if not query:
        return "Error: query must not be empty."
    max_results = min(int(args.get("max_results", 10)), 20)
    try:
        files = await asyncio.get_event_loop().run_in_executor(
            None, lambda: drive_search(query, max_results)
        )
    except Exception as exc:
        return f"Drive search failed: {exc}"
    if not files:
        return f"No Drive files found for: {query}"
    lines = [f"Drive files for '{query}' ({len(files)} found):\n"]
    for f in files:
        lines.append(
            f"ID: {f['id']}\nName: {f['name']}\nType: {f['type']}\n"
            f"Modified: {f['modified']}\nURL: {f['url']}"
        )
    return "\n\n".join(lines)


async def execute_list_drive_folder(ctx: ToolContext, args: dict[str, Any]) -> str:
    import asyncio
    from services.google_service import drive_list_folder, get_credentials
    if not get_credentials():
        return _google_not_connected()
    folder_id = str(args.get("folder_id", "root")).strip() or "root"
    drive_id = str(args.get("drive_id", "")).strip() or None
    max_results = min(int(args.get("max_results", 50)), 50)
    try:
        items = await asyncio.get_event_loop().run_in_executor(
            None, lambda: drive_list_folder(folder_id, drive_id=drive_id, max_results=max_results)
        )
    except Exception as exc:
        return f"Drive folder listing failed: {exc}"
    if not items:
        return f"Folder '{folder_id}' is empty or does not exist."
    lines = [f"Contents of Drive folder '{folder_id}' ({len(items)} items):\n"]
    for item in items:
        kind = "[FOLDER]" if item["type"] == "folder" else "[FILE]"
        lines.append(f"{kind} {item['name']}\n  ID: {item['id']}\n  URL: {item['url']}")
    return "\n\n".join(lines)


async def execute_list_shared_drives(ctx: ToolContext, args: dict[str, Any]) -> str:
    """List all shared (team) drives — the top-level folder trees beside My Drive."""
    import asyncio
    from services.google_service import drive_list_shared_drives, get_credentials
    if not get_credentials():
        return _google_not_connected()
    try:
        drives = await asyncio.get_event_loop().run_in_executor(
            None, drive_list_shared_drives
        )
    except Exception as exc:
        return f"Shared drive listing failed: {exc}"
    if not drives:
        return "No shared drives found. (Top-level folders may live in My Drive — use list_drive_folder with folder_id='root'.)"
    lines = [f"Shared drives ({len(drives)} found):\n"]
    for d in drives:
        lines.append(
            f"[SHARED DRIVE] {d['name']}\n  ID: {d['id']}\n"
            f"  To list its root folders: list_drive_folder with folder_id='{d['id']}' and drive_id='{d['id']}'"
        )
    return "\n\n".join(lines)


async def execute_search_drive_content(ctx: ToolContext, args: dict[str, Any]) -> str:
    """Search Drive and read the content of top matching files."""
    import asyncio
    from services.google_service import drive_search, drive_get_content, get_credentials
    if not get_credentials():
        return _google_not_connected()
    query = str(args.get("query", "")).strip()
    if not query:
        return "Error: query must not be empty."
    max_files = min(int(args.get("max_files", 3)), 5)

    # Step 1: search for matching files
    try:
        files = await asyncio.get_event_loop().run_in_executor(
            None, lambda: drive_search(query, max_files * 2)  # fetch extra, filter folders
        )
    except Exception as exc:
        return f"Drive search failed: {exc}"

    # Filter out folders and non-text types
    FOLDER_MIME = "application/vnd.google-apps.folder"
    READABLE = {
        "application/vnd.google-apps.document",
        "application/vnd.google-apps.spreadsheet",
        "application/vnd.google-apps.presentation",
        "text/plain", "text/csv", "text/markdown",
        "application/pdf",
    }
    readable_files = [
        f for f in files
        if f["type"] != FOLDER_MIME and (
            f["type"] in READABLE or f["type"].startswith("text/")
        )
    ][:max_files]

    if not readable_files:
        return f"No readable Drive files found for: {query}"

    # Step 2: read each file's content
    lines = [f"Drive content search for '{query}' ({len(readable_files)} file(s) read):\n"]
    for f in readable_files:
        try:
            result = await asyncio.get_event_loop().run_in_executor(
                None, lambda fid=f["id"]: drive_get_content(fid)
            )
            content = result.get("content", "").strip()
            if not content:
                lines.append(f"--- {f['name']} (no readable text) ---")
            else:
                # Cap each file at 4000 chars to avoid context overflow
                truncated = content[:4000]
                suffix = "..." if len(content) > 4000 else ""
                lines.append(
                    f"--- {result['name']} ({result['type']}) ---\n"
                    f"URL: {result['url']}\n\n"
                    f"{truncated}{suffix}"
                )
        except Exception as exc:
            lines.append(f"--- {f['name']} (read failed: {exc}) ---")

    return "\n\n".join(lines)


async def execute_read_drive_file(ctx: ToolContext, args: dict[str, Any]) -> str:
    import asyncio
    import re as _re
    from services.google_service import drive_get_content, get_credentials
    if not get_credentials():
        return _google_not_connected()
    file_id = str(args.get("file_id", "")).strip()
    # Accept full Docs/Drive/Sheets URLs — extract the file ID.
    m = _re.search(r"/d/([\w-]{20,})", file_id) or _re.search(r"[?&]id=([\w-]{20,})", file_id)
    if m:
        file_id = m.group(1)
    if not file_id:
        return "Error: file_id is required."
    try:
        offset = max(0, int(args.get("offset") or 0))
    except (TypeError, ValueError):
        offset = 0
    try:
        # Fetch the FULL text and slice a page locally — pagination lets the
        # model walk arbitrarily long contracts/specs instead of only ever
        # seeing the first page (field report: an Article-IV IP review was cut
        # mid-sentence with no way to page further).
        result = await asyncio.get_event_loop().run_in_executor(
            None, lambda: drive_get_content(file_id, max_chars=None)
        )
    except Exception as exc:
        return f"Failed to read Drive file: {exc}"
    content = result.get("content", "")
    if not content:
        return f"File '{result.get('name', file_id)}' has no readable text content."
    total = len(content)
    if offset >= total:
        return (
            f"Offset {offset:,} is past the end of '{result.get('name', file_id)}' "
            f"({total:,} characters total). The document has been fully read."
        )
    page = content[offset : offset + 30_000]
    end = offset + len(page)
    if end < total:
        note = (
            f"\n\n[PAGE ENDS at character {end:,} of {total:,} — the document CONTINUES. "
            f"Call read_drive_file again with offset={end} to read the next page. Do not "
            "draw final conclusions until you have read to the end.]"
        )
    else:
        note = (
            f"\n\n[END OF DOCUMENT — characters {offset:,}–{end:,} of {total:,}.]"
            if offset > 0
            else ""
        )
    return (
        f"File: {result['name']}\nType: {result['type']}\nURL: {result['url']}\n"
        f"Characters {offset:,}–{end:,} of {total:,}:\n\n"
        f"{page}{note}"
    )


async def execute_list_recent_drive_files(ctx: ToolContext, args: dict[str, Any]) -> str:
    import asyncio
    from services.google_service import drive_recent_files, get_credentials
    if not get_credentials():
        return _google_not_connected()
    max_results = min(int(args.get("max_results", 8) or 8), 20)
    try:
        files = await asyncio.get_event_loop().run_in_executor(
            None, lambda: drive_recent_files(max_results)
        )
    except Exception as exc:
        return f"Failed to list recent Drive files: {exc}"
    if not files:
        return "No recent Drive files found."
    lines = ["Most recently modified Drive files (newest first):"]
    for f in files:
        lines.append(
            f"- {f['name']} (modified {f['modified']}, owner {f['owner'] or 'unknown'}, "
            f"id={f['id']})"
        )
    lines.append(
        "Ask the user to confirm which one they mean, then call follow_drive_document."
    )
    return "\n".join(lines)


async def execute_follow_drive_document(ctx: ToolContext, args: dict[str, Any]) -> str:
    import asyncio
    from services.google_service import drive_get_content, get_credentials
    from services.live_document import extract_drive_file_id, set_followed_doc
    if not get_credentials():
        return _google_not_connected()
    file_id = extract_drive_file_id(str(args.get("file_id", "")))
    if not file_id:
        return "Error: file_id (or a pasted document URL) is required."
    try:
        result = await asyncio.get_event_loop().run_in_executor(
            None, lambda: drive_get_content(file_id)
        )
    except Exception as exc:
        return f"Could not read that document: {exc}"
    name = result.get("name") or "document"
    url = result.get("url") or ""
    await set_followed_doc(ctx.db, ctx.conversation_id, file_id, name, url)
    from services.workroom_context import auto_pin_if_room
    await auto_pin_if_room(ctx.db, ctx.conversation_id, "drive_doc", name, file_id)
    content = (result.get("content") or "").strip()
    preview = content[:2000] + ("…" if len(content) > 2000 else "")
    return (
        f"Now following \"{name}\" in this conversation — its current contents will be "
        "re-read automatically on every message, so you always see the user's latest "
        f"edits. Current text begins:\n---\n{preview}\n---"
    )


async def execute_unfollow_drive_document(ctx: ToolContext, _args: dict[str, Any]) -> str:
    from services.live_document import clear_followed_doc
    removed = await clear_followed_doc(ctx.db, ctx.conversation_id)
    return (
        "Stopped following the live document in this conversation."
        if removed
        else "No document was being followed in this conversation."
    )


async def execute_add_to_knowledge_base(ctx: ToolContext, args: dict[str, Any]) -> str:
    """Import a Drive file or a generated file into the Knowledge Base."""
    from repositories.document_repo import DocumentCategoryRepository
    from services.documents.ingestion import DocumentIngestionService, DuplicateDocumentError

    drive_ref = str(args.get("drive_file_id", "")).strip()
    gen_name = str(args.get("generated_filename", "")).strip()
    title = str(args.get("title", "")).strip() or None
    category = str(args.get("category", "")).strip()

    if not drive_ref and not gen_name:
        return (
            "Error: provide drive_file_id (a Drive ID or pasted URL) or "
            "generated_filename. Use search_drive / list_recent_drive_files to find "
            "Drive files, or list the Generated Files."
        )

    category_id = None
    if category:
        cat = await DocumentCategoryRepository(ctx.db).get_or_create(category)
        category_id = cat.id

    try:
        if drive_ref:
            from services.google_service import get_credentials
            from services.live_document import extract_drive_file_id
            from services.documents.drive_import import import_drive_file

            if not get_credentials():
                return _google_not_connected()
            doc = await import_drive_file(
                db=ctx.db,
                embedding_svc=ctx.embedding_service,
                file_id=extract_drive_file_id(drive_ref),
                title=title,
                category_id=category_id,
                is_regulated=False,
                created_by_id=ctx.user_id,
            )
        else:
            path = (_GENERATED_FILES_DIR / gen_name).resolve()
            if not str(path).startswith(str(_GENERATED_FILES_DIR.resolve())) or not path.is_file():
                return f"Error: generated file '{gen_name}' not found."
            svc = DocumentIngestionService(db=ctx.db, embedding_svc=ctx.embedding_service)
            doc = await svc.ingest(
                filename=path.name,
                raw_bytes=path.read_bytes(),
                title=title or path.stem,
                category_id=category_id,
                is_regulated=False,
                created_by_id=ctx.user_id,
            )
    except DuplicateDocumentError as exc:
        return (
            f"Already in the Knowledge Base: an identical document exists as "
            f"\"{exc.existing.title}\" — nothing was imported."
        )
    except ValueError as exc:
        return f"Could not import that file into the Knowledge Base: {exc}"
    except Exception as exc:  # noqa: BLE001
        return f"Knowledge Base import failed: {exc}"

    from services.workroom_context import auto_pin_if_room, log_room_event
    await auto_pin_if_room(ctx.db, ctx.conversation_id, "kb_doc", doc.title, str(doc.id))
    await log_room_event(
        ctx.db, ctx.conversation_id, f'Imported "{doc.title}" into the Knowledge Base'
    )

    return (
        f"Added to the Knowledge Base: \"{doc.title}\" [id={doc.id}"
        + (f", category={category}" if category else "")
        + "]. It is now searchable and citable."
    )


async def execute_get_file_template(ctx: ToolContext, args: dict[str, Any]) -> str:
    """Fetch the required structure for a document type from the templates doc."""
    from services.file_templates import get_file_template

    file_type = str(args.get("file_type", "")).strip()
    if not file_type:
        return "Error: file_type is required (e.g. 'memo', 'SOP', 'letter')."
    return await get_file_template(ctx.db, file_type)


async def execute_check_drive_backup_status(ctx: ToolContext, _args: dict[str, Any]) -> str:
    """Report whether the nightly Drive → GCS backup is current."""
    from services.backup_monitor import get_backup_status

    return await get_backup_status(ctx.db)


# ── Workroom tools ────────────────────────────────────────────────────────────

_WORKROOM_KINDS = (
    "drive_doc", "kb_doc", "generated_file", "note",
    "email_thread", "task", "odoo_record", "regulatory_doc",
)


async def _resolve_room_or_error(ctx: ToolContext, args: dict[str, Any]):
    """Shared resolution: (room, None) on success, (None, error_message) otherwise."""
    from services.workroom_context import resolve_workroom

    title_hint = str(args.get("workroom_title", "")).strip()
    room, titles = await resolve_workroom(
        ctx.db, ctx.user_id, ctx.conversation_id, title_hint
    )
    if room is not None:
        return room, None
    if not titles:
        return None, (
            "No workrooms exist yet. The user can create one from the Workrooms "
            "page (satellite next to the sun)."
        )
    listing = "; ".join(f'"{t}"' for t in titles)
    if title_hint:
        return None, (
            f'No workroom matches "{title_hint}". Active rooms: {listing}. '
            "Call again with one of these titles."
        )
    return None, (
        "This conversation is not inside a workroom. Pass workroom_title to "
        f"target one of the active rooms: {listing}."
    )


async def execute_create_workroom(ctx: ToolContext, args: dict[str, Any]) -> str:
    from models.db.workroom import Workroom
    from repositories.conversation_repo import ConversationRepository
    from services.workroom_context import add_journal_entry, list_active_workrooms

    title = str(args.get("title", "")).strip()[:200]
    goal = str(args.get("goal", "")).strip()[:4000]
    if not title:
        return "Error: title is required."
    # Guard against accidental duplicates of an existing active room.
    for r in await list_active_workrooms(ctx.db, ctx.user_id):
        if r.title.lower() == title.lower():
            return (
                f'A workroom named "{r.title}" already exists. Use it directly '
                "(add_to_workroom / list_workroom_items) or pick another title."
            )
    conv = await ConversationRepository(ctx.db).create(
        user_id=ctx.user_id, title=f"Workroom: {title}"
    )
    room = Workroom(
        user_id=ctx.user_id,
        title=title,
        goal=goal,
        conversation_id=conv.id,
    )
    ctx.db.add(room)
    await ctx.db.flush()
    await add_journal_entry(ctx.db, room, "Room created by Gerry in chat")
    return (
        f'Workroom "{title}" created'
        + (f" with goal: {goal}" if goal else "")
        + ". It has its own room conversation (open it from the Workrooms page "
        "or the chat sidebar). You can now pin items with add_to_workroom "
        f'(workroom_title="{title}") and log progress with log_workroom_progress.'
    )


async def execute_add_to_workroom(ctx: ToolContext, args: dict[str, Any]) -> str:
    from services.workroom_context import pin_workroom_item

    kind = str(args.get("kind", "")).strip().lower()
    label = str(args.get("label", "")).strip()
    ref_id = str(args.get("ref_id", "")).strip()
    if kind not in _WORKROOM_KINDS:
        return f"Error: kind must be one of: {', '.join(_WORKROOM_KINDS)}."
    if not label:
        return "Error: label is required."
    room, err = await _resolve_room_or_error(ctx, args)
    if err:
        return err
    item, created = await pin_workroom_item(ctx.db, room, kind, label, ref_id)
    if not created:
        return f'Already pinned in "{room.title}": [{kind}] {item.label}.'
    return (
        f'Pinned to workroom "{room.title}": [{kind}] {label}'
        + (f" (ref: {ref_id})" if ref_id else "")
        + ". It will appear in the room's context every turn."
    )


async def execute_list_workroom_items(ctx: ToolContext, args: dict[str, Any]) -> str:
    from sqlalchemy import desc as _desc, select as _select

    from models.db.workroom import WorkroomItem, WorkroomJournalEntry

    room, err = await _resolve_room_or_error(ctx, args)
    if err:
        return err
    items = list(
        (
            await ctx.db.execute(
                _select(WorkroomItem)
                .where(WorkroomItem.workroom_id == room.id)
                .order_by(WorkroomItem.created_at)
            )
        ).scalars()
    )
    journal = list(
        (
            await ctx.db.execute(
                _select(WorkroomJournalEntry)
                .where(WorkroomJournalEntry.workroom_id == room.id)
                .order_by(_desc(WorkroomJournalEntry.created_at))
                .limit(10)
            )
        ).scalars()
    )
    lines = [f'Workroom "{room.title}"']
    if room.goal.strip():
        lines.append(f"Goal: {room.goal.strip()}")
    if items:
        lines.append(f"Pinned items ({len(items)}):")
        for it in items:
            ref = f" (ref: {it.ref_id})" if it.ref_id else ""
            lines.append(f"- [{it.kind}] {it.label}{ref}")
    else:
        lines.append("No items pinned yet.")
    if journal:
        lines.append("Recent journal (newest first):")
        for j in journal:
            stamp = j.created_at.strftime("%Y-%m-%d") if j.created_at else ""
            lines.append(f"- {stamp}: {j.entry}")
    return "\n".join(lines)


async def execute_log_workroom_progress(ctx: ToolContext, args: dict[str, Any]) -> str:
    from services.workroom_context import add_journal_entry

    entry = str(args.get("entry", "")).strip()
    if not entry:
        return "Error: entry is required — one sentence describing the progress."
    room, err = await _resolve_room_or_error(ctx, args)
    if err:
        return err
    await add_journal_entry(ctx.db, room, entry)
    return f'Logged to "{room.title}" journal: {entry}'


async def execute_get_calendar_events(ctx: ToolContext, args: dict[str, Any]) -> str:
    import asyncio
    from services.google_service import calendar_events, get_credentials
    if not get_credentials():
        return _google_not_connected()
    days_behind = int(args.get("days_behind", 0))
    days_ahead = int(args.get("days_ahead", 7))
    try:
        events = await asyncio.get_event_loop().run_in_executor(
            None, lambda: calendar_events(days_behind, days_ahead)
        )
    except Exception as exc:
        return f"Calendar fetch failed: {exc}"
    if not events:
        return f"No calendar events in the next {days_ahead} day(s)."
    lines = [f"Calendar events ({len(events)} found):\n"]
    for e in events:
        attendee_str = ", ".join(e["attendees"]) if e["attendees"] else ""
        lines.append(
            f"• {e['title']}\n"
            f"  Start: {e['start']}  End: {e['end']}\n"
            + (f"  Location: {e['location']}\n" if e["location"] else "")
            + (f"  Attendees: {attendee_str}\n" if attendee_str else "")
            + (f"  {e['description']}\n" if e["description"] else "")
            + f"  URL: {e['url']}"
        )
    return "\n\n".join(lines)


async def execute_add_contacts(ctx: ToolContext, args: dict[str, Any]) -> str:
    """Upsert manual contacts into Gerry's own contact store (the app's
    Communications → Contacts page). Mirrors the POST /api/google/contacts
    endpoint's shape so the page shows them identically."""
    from services import email_contacts as ec

    raw = args.get("contacts")
    if isinstance(raw, dict):
        raw = [raw]
    if not isinstance(raw, list) or not raw:
        return "Error: 'contacts' must be a non-empty array of {email, name?, company?, notes?}."

    contacts = await ec.get_contacts(ctx.db)
    added: list[str] = []
    updated: list[str] = []
    skipped: list[str] = []
    for item in raw:
        if not isinstance(item, dict):
            skipped.append(str(item)[:60])
            continue
        email = ec.extract_email(str(item.get("email", "")))
        if not email or "@" not in email:
            skipped.append(str(item.get("email", "(no email)"))[:60])
            continue
        domain = ec.domain_of(email)
        existing = contacts.get(email, {})
        name = str(item.get("name") or "").strip() or (existing.get("name") or "")
        company = (
            str(item.get("company") or "").strip()
            or existing.get("company")
            or ec.company_from_domain(domain)
        )
        notes = str(item.get("notes") or "").strip() or (existing.get("notes") or "")
        contacts[email] = {
            "email": email,
            "name": name,
            "company": company,
            "domain": domain,
            "notes": notes,
            "source": "manual",
            "count": int(existing.get("count", 0) or 0),
            "last_seen": existing.get("last_seen", ""),
        }
        (updated if existing else added).append(f"{name or email} <{email}>")

    if added or updated:
        await ec.save_contacts(ctx.db, contacts, ctx.user_id)

    parts: list[str] = []
    if added:
        parts.append(f"Added {len(added)} contact(s): " + "; ".join(added))
    if updated:
        parts.append(f"Updated {len(updated)} existing contact(s): " + "; ".join(updated))
    if skipped:
        parts.append(f"Skipped {len(skipped)} entry(ies) without a valid email: " + "; ".join(skipped))
    if not parts:
        return "No contacts were added — every entry was missing a valid email address."
    parts.append("They now appear on the Contacts page (Communications → Contacts).")
    return "\n".join(parts)


async def execute_search_contacts(ctx: ToolContext, args: dict[str, Any]) -> str:
    import asyncio
    from services.email_contacts import get_contacts, search_contacts_store
    from services.google_service import contacts_search, get_credentials

    query = str(args.get("query", "")).strip()
    if not query:
        return "Error: query must not be empty."

    # Gerry's own contacts, derived from email senders + manual entries.
    local: list[dict] = []
    try:
        store = await get_contacts(ctx.db)
        local = search_contacts_store(store, query, 10)
    except Exception:
        local = []

    # Google Contacts (only if the account is connected).
    google: list[dict] = []
    if get_credentials():
        try:
            google = await asyncio.get_event_loop().run_in_executor(
                None, lambda: contacts_search(query, 10)
            )
        except Exception:
            google = []

    if not local and not google:
        return f"No contacts found for: {query}"

    lines: list[str] = []
    if local:
        lines.append(f"PMI contacts (from email) matching '{query}':")
        for c in local:
            company = f" — {c['company']}" if c.get("company") else ""
            seen = f" (seen {c.get('count', 0)}x)" if c.get("count") else ""
            lines.append(
                f"Name: {c.get('name') or '(unknown)'}{company}\nEmail: {c['email']}{seen}"
            )
    if google:
        lines.append(f"Google Contacts matching '{query}':")
        for c in google:
            lines.append(
                f"Name: {c['name']}\nEmail: {c['email']}\n"
                f"Phone: {c['phone']}\nCompany: {c['company']}"
            )
    return "\n\n".join(lines)


async def execute_read_google_sheet(ctx: ToolContext, args: dict[str, Any]) -> str:
    import asyncio
    from services.google_service import sheets_read, get_credentials
    if not get_credentials():
        return _google_not_connected()
    spreadsheet_id = str(args.get("spreadsheet_id", "")).strip()
    if not spreadsheet_id:
        return "Error: spreadsheet_id is required."
    range_ = str(args.get("range", "Sheet1")).strip() or "Sheet1"
    try:
        result = await asyncio.get_event_loop().run_in_executor(
            None, lambda: sheets_read(spreadsheet_id, range_)
        )
    except Exception as exc:
        return f"Failed to read spreadsheet: {exc}"
    rows = result.get("rows", [])
    if not rows:
        return f"Spreadsheet range '{range_}' is empty."
    # Format as a simple table
    lines = [f"Sheet data ({result['row_count']} rows, range: {result['range']}):\n"]
    for row in rows[:100]:  # cap at 100 rows
        lines.append("\t".join(str(cell) for cell in row))
    if result["row_count"] > 100:
        lines.append(f"... ({result['row_count'] - 100} more rows)")
    return "\n".join(lines)


async def execute_list_google_tasks(ctx: ToolContext, args: dict[str, Any]) -> str:
    import asyncio
    from services.google_service import tasks_list, get_credentials
    if not get_credentials():
        return _google_not_connected()
    max_results = min(int(args.get("max_results", 25)), 50)
    try:
        tasks = await asyncio.get_event_loop().run_in_executor(
            None, lambda: tasks_list(max_results)
        )
    except Exception as exc:
        return f"Failed to fetch Google Tasks: {exc}"
    if not tasks:
        return "No incomplete Google Tasks found."
    lines = [f"Google Tasks ({len(tasks)} incomplete):\n"]
    for t in tasks:
        due_str = f", due {t['due'][:10]}" if t.get("due") else ""
        notes_str = f"\n  Notes: {t['notes']}" if t.get("notes") else ""
        lines.append(f"• [{t['list']}] {t['title']}{due_str}{notes_str}")
    return "\n".join(lines)


_GENERATED_FILES_DIR = Path(__file__).resolve().parent.parent.parent / "generated_files"


async def execute_generate_file(ctx: ToolContext, args: dict[str, Any]) -> str:
    filename = str(args.get("filename", "output.txt")).strip()
    content  = str(args.get("content", ""))

    # Sanitize filename
    filename = re.sub(r"[^\w.\-]", "_", filename)
    allowed = {".txt", ".md", ".csv", ".json"}
    suffix = Path(filename).suffix.lower()
    if suffix not in allowed:
        filename = Path(filename).stem + ".txt"

    _GENERATED_FILES_DIR.mkdir(exist_ok=True)
    uid_prefix = uuid.uuid4().hex[:8]
    safe_name = f"{uid_prefix}_{filename}"
    out_path = _GENERATED_FILES_DIR / safe_name
    out_path.write_text(content, encoding="utf-8")

    # Read-back verification: confirm the file actually landed on disk before
    # reporting success, so a silent write failure can never be reported as done.
    if not out_path.is_file():
        return "Error: file was not created — nothing was written to disk. Treat as failed."
    written = out_path.stat().st_size
    if content and written == 0:
        return "Error: file was created empty despite having content. Treat as failed."
    from services.workroom_context import auto_pin_if_room, log_room_event
    await auto_pin_if_room(ctx.db, ctx.conversation_id, "generated_file", filename, safe_name)
    await log_room_event(ctx.db, ctx.conversation_id, f"Gerry generated file: {filename}")
    return f"File created and verified ({written} bytes): /api/files/{safe_name}"


def _add_markdown_runs(paragraph, text: str) -> None:
    """Add text to a python-docx paragraph, rendering **bold** spans as bold runs."""
    for i, segment in enumerate(re.split(r"\*\*(.+?)\*\*", text)):
        if not segment:
            continue
        run = paragraph.add_run(segment)
        run.bold = bool(i % 2)  # odd indices are the captured bold groups


async def execute_create_docx(ctx: ToolContext, args: dict[str, Any]) -> str:
    """Build a real Word (.docx) document from lightweight Markdown content.

    Layout support (driven by company templates / style guide): base font and
    size, real page headers/footers (with automatic 'Page X of Y' fields), and
    pipe tables rendered as Word tables with a colored header row and
    alternating row shading — so template rules like "navy header row" or
    "footer: CONFIDENTIAL" actually appear in the file, not just in the text.
    """
    import docx
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    filename = str(args.get("filename", "document.docx")).strip()
    title = str(args.get("title", "")).strip()
    content = str(args.get("content", ""))
    font_name = str(args.get("font", "")).strip()
    try:
        font_size = float(args.get("font_size") or 0)
    except (TypeError, ValueError):
        font_size = 0
    header_left = str(args.get("header_left", "")).strip()
    header_right = str(args.get("header_right", "")).strip()
    footer_left = str(args.get("footer_left", "")).strip()
    accent = re.sub(r"[^0-9A-Fa-f]", "", str(args.get("accent_color", "")))[:6] or "1F3864"
    banner_label = str(args.get("banner_label", "")).strip()
    banner_title = str(args.get("banner_title", "")).strip()
    banner_subtitle = str(args.get("banner_subtitle", "")).strip()
    # Light tint of the accent for banner secondary lines (label/subtitle).
    _a = [int(accent[i : i + 2], 16) for i in (0, 2, 4)]
    tint = "".join(f"{c + int((255 - c) * 0.62):02X}" for c in _a)

    # Sanitize filename and force a .docx extension
    filename = re.sub(r"[^\w.\- ]", "_", filename).strip() or "document.docx"
    if Path(filename).suffix.lower() not in (".docx", ".doc"):
        filename = Path(filename).stem + ".docx"
    if Path(filename).suffix.lower() == ".doc":
        # python-docx only writes the modern format
        filename = Path(filename).stem + ".docx"

    document = docx.Document()

    # Base typography (applies to Normal; headings keep their theme fonts).
    if font_name:
        document.styles["Normal"].font.name = font_name
    if font_size > 0:
        document.styles["Normal"].font.size = Pt(font_size)
    # Brand the heading styles with the accent color (reference docs use the
    # company navy/teal for headings rather than Word's default blue).
    for h in ("Heading 1", "Heading 2", "Heading 3"):
        try:
            document.styles[h].font.color.rgb = RGBColor.from_string(accent)
            if font_name:
                document.styles[h].font.name = font_name
        except Exception:  # noqa: BLE001 — style may be absent in exotic bases
            pass

    def _page_field(paragraph, code: str) -> None:
        """Append a Word field (PAGE / NUMPAGES) to a paragraph."""
        run = paragraph.add_run()
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = f" {code} "
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        run._r.append(begin)
        run._r.append(instr)
        run._r.append(end)

    # Page header/footer — the built-in Header/Footer styles carry centre +
    # right tab stops, so "left\t\tright" lands text on both edges. Reference
    # documents render these in small gray type: 9pt, #666666.
    def _hf_style(paragraph) -> None:
        for r in paragraph.runs:
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor.from_string("666666")
            if font_name:
                r.font.name = font_name

    section = document.sections[0]
    if header_left or header_right:
        hp = section.header.paragraphs[0]
        hp.text = ""
        hp.add_run(f"{header_left}\t\t{header_right}")
        _hf_style(hp)
    if footer_left or content:
        fp = section.footer.paragraphs[0]
        fp.text = ""
        if footer_left:
            fp.add_run(f"{footer_left}\t\tPage ")
        else:
            fp.add_run("\t\tPage ")
        _page_field(fp, "PAGE")
        fp.add_run(" of ")
        _page_field(fp, "NUMPAGES")
        _hf_style(fp)

    def _shade(cell, fill: str) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), fill)
        tc_pr.append(shd)

    def _thin_borders(table) -> None:
        """Thin single borders on every edge — the label/value grid look."""
        tbl_pr = table._tbl.tblPr
        borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            el = OxmlElement(f"w:{edge}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "4")
            el.set(qn("w:color"), "404040")
            borders.append(el)
        tbl_pr.append(borders)

    if banner_label or banner_title or banner_subtitle:
        # Title-block banner: one full-width accent-filled cell with centred
        # label (small, tinted) / title (large, white) / subtitle (small, tinted)
        # — matches the company template's cover banner.
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        bt = document.add_table(rows=1, cols=1)
        bcell = bt.rows[0].cells[0]
        _shade(bcell, accent)
        first = True
        for text, size, bold, color in (
            (banner_label, 9, True, tint),
            (banner_title, 22, True, "FFFFFF"),
            (banner_subtitle, 10, False, tint),
        ):
            if not text:
                continue
            bp = bcell.paragraphs[0] if first else bcell.add_paragraph()
            first = False
            bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            br = bp.add_run(text)
            br.bold = bold
            br.font.size = Pt(size)
            br.font.color.rgb = RGBColor.from_string(color)
        document.add_paragraph()

    def _flush_table(rows: list[list[str]], has_header: bool) -> None:
        if not rows:
            return
        cols = max(len(r) for r in rows)
        table = document.add_table(rows=len(rows), cols=cols)
        if not has_header:
            _thin_borders(table)
        for ri, row in enumerate(rows):
            for ci in range(cols):
                cell = table.rows[ri].cells[ci]
                text = row[ci] if ci < len(row) else ""
                para = cell.paragraphs[0]
                _add_markdown_runs(para, text)
                if has_header and ri == 0:
                    _shade(cell, accent)
                    for r in para.runs:
                        r.bold = True
                        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                elif has_header and ri % 2 == 0:
                    _shade(cell, "F2F2F2")
                elif not has_header and ci % 2 == 0:
                    # Label/value grid: label COLUMNS (1st, 3rd, ...) carry the
                    # accent fill with bold white text; value columns stay plain.
                    _shade(cell, accent)
                    for r in para.runs:
                        r.bold = True
                        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    if title:
        document.add_heading(title, level=0)

    table_rows: list[list[str]] = []
    table_has_header = False
    for raw_line in content.splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()
        # Pipe-table rows are collected and flushed as one Word table. A
        # markdown separator row ('| --- | --- |') after the first row marks it
        # as a HEADER table (colored first row); tables without a separator
        # render as plain label/value grids with light alternating shading.
        if stripped.startswith("|") and stripped.endswith("|") and len(stripped) > 1:
            cells = [c.strip() for c in stripped.strip("|").split("|")]
            if all(re.fullmatch(r":?-{2,}:?", c or "---") for c in cells):
                if len(table_rows) == 1:
                    table_has_header = True
                continue  # markdown separator row
            table_rows.append(cells)
            continue
        _flush_table(table_rows, table_has_header)
        table_rows = []
        table_has_header = False
        if not stripped:
            continue
        if stripped.startswith("### "):
            document.add_heading(stripped[4:].strip(), level=3)
        elif stripped.startswith("## "):
            document.add_heading(stripped[3:].strip(), level=2)
        elif stripped.startswith("# "):
            document.add_heading(stripped[2:].strip(), level=1)
        elif stripped.startswith(("- ", "* ")):
            p = document.add_paragraph(style="List Bullet")
            _add_markdown_runs(p, stripped[2:].strip())
        elif re.match(r"^\d+\.\s+", stripped):
            p = document.add_paragraph(style="List Number")
            _add_markdown_runs(p, re.sub(r"^\d+\.\s+", "", stripped))
        else:
            p = document.add_paragraph()
            _add_markdown_runs(p, stripped)
    _flush_table(table_rows, table_has_header)

    _GENERATED_FILES_DIR.mkdir(exist_ok=True)
    uid_prefix = uuid.uuid4().hex[:8]
    safe_name = f"{uid_prefix}_{filename}"
    out_path = _GENERATED_FILES_DIR / safe_name
    document.save(str(out_path))

    # Read-back verification: confirm the .docx exists and is a valid, openable
    # document before reporting success — never claim a document that isn't real.
    if not out_path.is_file() or out_path.stat().st_size == 0:
        return "Error: Word document was not created on disk. Treat as failed."
    try:
        docx.Document(str(out_path))  # re-open to prove it's a valid .docx
    except Exception as exc:  # noqa: BLE001 — report a corrupt write honestly
        return f"Error: Word document was written but is not a valid file ({exc}). Treat as failed."
    from services.workroom_context import auto_pin_if_room, log_room_event
    await auto_pin_if_room(ctx.db, ctx.conversation_id, "generated_file", filename, safe_name)
    await log_room_event(ctx.db, ctx.conversation_id, f"Gerry created Word document: {filename}")
    return (
        f"Word document created and verified ({out_path.stat().st_size} bytes): "
        f"/api/files/{safe_name}"
    )


async def execute_upload_to_drive(ctx: ToolContext, args: dict[str, Any]) -> str:
    """Upload a generated file to the user's Google Drive."""
    import asyncio

    from services.google_service import drive_get_metadata, drive_upload_file

    filename = str(args.get("filename", "")).strip()
    if not filename:
        return "Error: filename is required."
    # Guard against path traversal — only allow a bare filename.
    if "/" in filename or "\\" in filename or filename.startswith(".."):
        return "Error: invalid filename."

    path = (_GENERATED_FILES_DIR / filename).resolve()
    if not str(path).startswith(str(_GENERATED_FILES_DIR.resolve())) or not path.is_file():
        return f"Error: file '{filename}' not found in Generated Files. Create it first, then upload."

    # Strip the internal 8-hex id prefix (e.g. 'a1b2c3d4_Name.docx') for a clean Drive name.
    default_name = re.sub(r"^[0-9a-f]{8}_", "", filename)
    drive_name = str(args.get("drive_name") or default_name).strip() or default_name
    folder_id = str(args.get("folder_id") or "").strip() or None

    try:
        result = await asyncio.get_event_loop().run_in_executor(
            None, lambda: drive_upload_file(str(path), name=drive_name, folder_id=folder_id)
        )
    except Exception as exc:  # noqa: BLE001 — surface a readable message to the model
        return f"Drive upload failed: {exc}"

    # Read-back verification: confirm Drive actually has the file by fetching its
    # metadata with the returned id. Without a real, confirmed id we never report
    # a successful upload (this is exactly what prevents fabricated Drive links).
    file_id = str(result.get("id", "")).strip()
    if not file_id:
        return "Drive upload failed: no file id was returned. Treat as failed."
    try:
        meta = await asyncio.get_event_loop().run_in_executor(
            None, lambda: drive_get_metadata(file_id)
        )
    except Exception as exc:  # noqa: BLE001 — report verification failure honestly
        return f"Drive upload could not be verified ({exc}). Treat as unconfirmed."
    if meta is None or meta.get("trashed"):
        return (
            "Drive upload could not be verified — the file is not retrievable on "
            "Drive after upload. Treat as failed; do not report it as uploaded."
        )

    return (
        f"Uploaded and verified '{meta.get('name', drive_name)}' on Google Drive "
        f"(id={file_id}). Link: {meta.get('url') or result.get('url') or '(no link returned)'}"
    )



# ── Dispatcher ────────────────────────────────────────────────────────────────

TOOL_EXECUTORS = {
    "search_knowledge_base": execute_search_knowledge_base,
    "read_knowledge_base_document": execute_read_knowledge_base_document,
    "request_kb_deletion": execute_request_kb_deletion,
    "create_task": execute_create_task,
    "create_email_draft": execute_create_email_draft,
    "request_approval": execute_request_approval,
    "propose_odoo_write": execute_propose_odoo_write,
    "get_pending_approvals": execute_get_pending_approvals,
    "get_tasks": execute_get_tasks,
    "get_regulatory_status": execute_get_regulatory_status,
    "search_web": execute_search_web,
    "fetch_page": execute_fetch_page,
    "generate_file": execute_generate_file,
    "create_docx": execute_create_docx,
    "upload_to_drive": execute_upload_to_drive,
    # Google Workspace (read-only)
    "search_gmail": execute_search_gmail,
    "read_gmail_message": execute_read_gmail_message,
    "search_drive": execute_search_drive,
    "search_drive_content": execute_search_drive_content,
    "list_drive_folder": execute_list_drive_folder,
    "list_shared_drives": execute_list_shared_drives,
    "read_drive_file": execute_read_drive_file,
    "list_recent_drive_files": execute_list_recent_drive_files,
    "follow_drive_document": execute_follow_drive_document,
    "unfollow_drive_document": execute_unfollow_drive_document,
    "add_to_knowledge_base": execute_add_to_knowledge_base,
    "check_drive_backup_status": execute_check_drive_backup_status,
    "get_file_template": execute_get_file_template,
    "create_workroom": execute_create_workroom,
    "add_to_workroom": execute_add_to_workroom,
    "list_workroom_items": execute_list_workroom_items,
    "log_workroom_progress": execute_log_workroom_progress,
    "get_calendar_events": execute_get_calendar_events,
    "search_contacts": execute_search_contacts,
    "add_contacts": execute_add_contacts,
    "read_google_sheet": execute_read_google_sheet,
    "list_google_tasks": execute_list_google_tasks,
}

# House Manager custodian tools (registered here so dispatch_tool can run them;
# only agents whose TOOLS whitelist includes them can actually call them).
from services.agent.custodian_tools import CUSTODIAN_EXECUTORS  # noqa: E402

TOOL_EXECUTORS.update(CUSTODIAN_EXECUTORS)


# When a model passes its single string argument as plain text instead of a
# JSON object, the v2 normalizer delivers it as {"input": "<text>"}. Map that
# onto each tool's primary parameter so e.g. search_knowledge_base("NAR contract")
# works instead of failing with "Error: query must not be empty."
_PRIMARY_ARG = {
    "search_knowledge_base": "query",
    "read_knowledge_base_document": "query",
    "request_kb_deletion": "query",
    "search_web": "query",
    "search_gmail": "query",
    "search_drive": "query",
    "search_drive_content": "query",
    "search_contacts": "query",
    "add_contacts": "contacts",
    "fetch_page": "url",
    "read_gmail_message": "message_id",
    "read_drive_file": "file_id",
    "list_recent_drive_files": "max_results",
    "follow_drive_document": "file_id",
    "add_to_knowledge_base": "drive_file_id",
    "get_file_template": "file_type",
    "create_workroom": "title",
    "add_to_workroom": "label",
    "list_workroom_items": "workroom_title",
    "log_workroom_progress": "entry",
    "list_drive_folder": "folder_id",
    "read_google_sheet": "spreadsheet_id",
}


async def dispatch_tool(ctx: ToolContext, name: str, args: dict[str, Any]) -> str:
    """Execute a tool by name and return a string result for the model."""
    executor = TOOL_EXECUTORS.get(name)
    if executor is None:
        return f"Unknown tool: {name}"
    primary = _PRIMARY_ARG.get(name)
    if primary and primary not in args and isinstance(args.get("input"), str):
        args = {**args, primary: args["input"]}
    try:
        result = await executor(ctx, args)
    except Exception as exc:
        logger.exception("Tool %s raised", name)
        return f"Tool '{name}' failed: {exc}"
    if isinstance(result, str) and result.startswith("Error:"):
        # WARNING so it reaches app.log (the file handler drops INFO)
        logger.warning("Tool %s returned error (args keys=%s): %s", name, sorted(args), result[:200])
    return result
