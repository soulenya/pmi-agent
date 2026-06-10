"""Render lightweight Markdown into Word (.docx) bytes.

Mirrors the renderer used by the agent's create_docx tool: headings,
bullet / numbered lists, and **bold** spans. Anything fancier stays as
plain paragraphs so the output is always a valid, editable document.
"""

from __future__ import annotations

import io
import re


def _add_markdown_runs(paragraph, text: str) -> None:
    """Add text to a python-docx paragraph, rendering **bold** spans as bold runs."""
    for i, segment in enumerate(re.split(r"\*\*(.+?)\*\*", text)):
        if not segment:
            continue
        run = paragraph.add_run(segment)
        run.bold = bool(i % 2)  # odd indices are the captured bold groups


def markdown_to_docx_bytes(title: str, content: str) -> bytes:
    """Build a .docx document from lightweight Markdown and return its bytes."""
    import docx

    document = docx.Document()
    if title:
        document.add_heading(title, level=0)

    for raw_line in content.splitlines():
        stripped = raw_line.strip()
        if not stripped:
            continue
        if stripped.startswith("#### "):
            document.add_heading(stripped[5:].strip(), level=4)
        elif stripped.startswith("### "):
            document.add_heading(stripped[4:].strip(), level=3)
        elif stripped.startswith("## "):
            document.add_heading(stripped[3:].strip(), level=2)
        elif stripped.startswith("# "):
            document.add_heading(stripped[2:].strip(), level=1)
        elif stripped.startswith(("- ", "* ")):
            p = document.add_paragraph(style="List Bullet")
            _add_markdown_runs(p, stripped[2:].strip())
        elif re.match(r"^\d+\.\s+", stripped):
            p = document.add_paragraph(style="List Number")
            _add_markdown_runs(p, re.sub(r"^\d+\.\s+", "", stripped))
        else:
            p = document.add_paragraph()
            _add_markdown_runs(p, stripped)

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()
